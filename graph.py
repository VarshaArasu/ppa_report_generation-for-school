import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
import io
items = ['III', 'IV', 'V', 'VI', 'VII', 'VIII']
initial_scores = [34.01, 35.53, 35.70, 43.53, 45.67, 26.48]  
post_scores = [52.90, 45.63, 48.56, 50.15, 54.31, 33.25]   
x = np.arange(len(items))  
width = 0.25  
fig, ax = plt.subplots(figsize=(9, 5))
plt.gcf().set_facecolor('darkblue')  
ax.set_facecolor('darkblue') 
gap = 0.02
bar1 = ax.bar(x - width/2 - gap, initial_scores, width, label='Initial Assessment', color='purple')  
bar2 = ax.bar(x + width/2 + gap, post_scores, width, label='Post Program Assessment', color='green') 
ax.set_title('BSPI growth-Initial vs Post Program Assessment', color='white')  
ax.set_xticks(x)
ax.set_xticklabels(items, color='white') 
ax.tick_params(axis='y', colors='white')  
ax.tick_params(axis='x', colors='white')  
ax.legend(facecolor='gray', edgecolor='white', loc='upper right', fontsize=10, framealpha=0.7)  
for i, rect in enumerate(bar1):
    height = rect.get_height()
    ax.text(rect.get_x() + rect.get_width() / 2, height, str(initial_scores[i]), ha='center', va='bottom', color='white')
for i, rect in enumerate(bar2):
    height = rect.get_height()
    ax.text(rect.get_x() + rect.get_width() / 2, height, str(post_scores[i]), ha='center', va='bottom', color='white')
plt.tight_layout()
img_stream = io.BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0)  
doc = Document('page1_doc.docx')  
doc.add_paragraph("Here is the graph comparing Initial vs Post Program Assessment:")
doc.add_picture(img_stream, width=Inches(5))  
doc.save('page1_doc.docx')
img_stream.close()
print("Graph successfully added to the Word document.")
