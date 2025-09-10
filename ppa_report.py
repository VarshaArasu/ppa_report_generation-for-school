# Empty Word Document
from docx import Document
doc = Document()
doc.save('ppa_report.docx')

#JSON 
import json
with open('data.json', 'r') as file:
    data = json.load(file) 
schoolname = data['schoolname']
programstartdate = data['programstartdate']
numberofstudentsenrolled = data['numberofstudentsenrolled']
noofstudentsdoneppa = data['noofstudentsdoneppa']
gradesconducted = data['gradesconducted']
reportingperiod = data['reportingperiod']
minutestrained = data['minutestrained']
puzzlesattempted = data['puzzlesattempted']
puzzlesolved = data['puzzlesolved']
bspiscorerangeIA_80 = data['bspiscorerangeIA_80']
bspiscorerangeIA_60 = data['bspiscorerangeIA_60']
bspiscorerangeIA_40 = data['bspiscorerangeIA_40']
bspiscorerangeIA_20 = data['bspiscorerangeIA_20']
bspiscorerangeIA_10 = data['bspiscorerangeIA_10']
bspiscorerangePPA_80 = data['bspiscorerangePPA_80']
bspiscorerangePPA_60 = data['bspiscorerangePPA_60']
bspiscorerangePPA_40 = data['bspiscorerangePPA_40']
bspiscorerangePPA_20 = data['bspiscorerangePPA_20']
bspiscorerangePPA_10 = data['bspiscorerangePPA_10']
noofstudIA = data['noofstudIA']
noofstudPPA = data['noofstudPPA']
schoolavgbspi = data['schoolavgbspi']
bspigrowthgraphIA1 = data['bspigrowthgraphIA1']
bspigrowthgraphIA2 = data['bspigrowthgraphIA2']
bspigrowthgraphIA3 = data['bspigrowthgraphIA3']
bspigrowthgraphIA4 = data['bspigrowthgraphIA4']
bspigrowthgraphIA5 = data['bspigrowthgraphIA5']
bspigrowthgraphIA6 = data['bspigrowthgraphIA6']
bspigrowthgraphPPA1 = data['bspigrowthgraphPPA1']
bspigrowthgraphPPA2 = data['bspigrowthgraphPPA2']
bspigrowthgraphPPA3 = data['bspigrowthgraphPPA3']
bspigrowthgraphPPA4 = data['bspigrowthgraphPPA4']
bspigrowthgraphPPA5 = data['bspigrowthgraphPPA5']
bspigrowthgraphPPA6 = data['bspigrowthgraphPPA6']

grade3IA20 = data['grade3IA20']
grade3IA40 = data['grade3IA40']
grade3IA60 = data['grade3IA60']
grade3IA80 = data['grade3IA80']
grade3IA100 = data['grade3IA100']
grade3PPA20 = data['grade3PPA20']
grade3PPA40 = data['grade3PPA40']
grade3PPA60 = data['grade3PPA60']
grade3PPA80 = data['grade3PPA80']
grade3PPA100 = data['grade3PPA100']

grade4IA20 = data['grade4IA20']
grade4IA40 = data['grade4IA40']
grade4IA60 = data['grade4IA60']
grade4IA80 = data['grade4IA80']
grade4IA100 = data['grade4IA100']
grade4PPA20 = data['grade4PPA20']
grade4PPA40 = data['grade4PPA40']
grade4PPA60 = data['grade4PPA60']
grade4PPA80 = data['grade4PPA80']
grade4PPA100 = data['grade4PPA100']

grade5IA20 = data['grade5IA20']
grade5IA40 = data['grade5IA40']
grade5IA60 = data['grade5IA60']
grade5IA80 = data['grade5IA80']
grade5IA100 = data['grade5IA100']
grade5PPA20 = data['grade5PPA20']
grade5PPA40 = data['grade5PPA40']
grade5PPA60 = data['grade5PPA60']
grade5PPA80 = data['grade5PPA80']
grade5PPA100 = data['grade5PPA100']

grade6IA20 = data['grade6IA20']
grade6IA40 = data['grade6IA40']
grade6IA60 = data['grade6IA60']
grade6IA80 = data['grade6IA80']
grade6IA100 = data['grade6IA100']
grade6PPA20 = data['grade6PPA20']
grade6PPA40 = data['grade6PPA40']
grade6PPA60 = data['grade6PPA60']
grade6PPA80 = data['grade6PPA80']
grade6PPA100 = data['grade6PPA100']

grade7IA20 = data['grade7IA20']
grade7IA40 = data['grade7IA40']
grade7IA60 = data['grade7IA60']
grade7IA80 = data['grade7IA80']
grade7IA100 = data['grade7IA100']
grade7PPA20 = data['grade7PPA20']
grade7PPA40 = data['grade7PPA40']
grade7PPA60 = data['grade7PPA60']
grade7PPA80 = data['grade7PPA80']
grade7PPA100 = data['grade7PPA100']

grade8IA20 = data['grade8IA20']
grade8IA40 = data['grade8IA40']
grade8IA60 = data['grade8IA60']
grade8IA80 = data['grade8IA80']
grade8IA100 = data['grade8IA100']
grade8PPA20 = data['grade8PPA20']
grade8PPA40 = data['grade8PPA40']
grade8PPA60 = data['grade8PPA60']
grade8PPA80 = data['grade8PPA80']
grade8PPA100 = data['grade8PPA100']

highly_prodigiesname = data['highly_prodigiesname']
highly_prodigiesgrade = data['highly_prodigiesgrade']
highly_prodigiessec = data['highly_prodigiessec']
highly_prodigiesBSPI = data['highly_prodigiesBSPI']

FAstud1 = data['FAstud1']
FAgrade1 = data['FAgrade1']
FAsec1 = data['FAsec1']
FAppa_bspi1 = data['FAppa_bspi1']
FAstud2 = data['FAstud2']
FAgrade2 = data['FAgrade2']
FAsec2 = data['FAsec2']
FAppa_bspi2 = data['FAppa_bspi2']
FAstud3 = data['FAstud3']
FAgrade3 = data['FAgrade3']
FAsec3 = data['FAsec3']
FAppa_bspi3 = data['FAppa_bspi3']

gstars_stud1 = data['gstars_stud1']
gstars_grade1 = data['gstars_grade1']
gstars_sec1 = data['gstars_sec1']
gstars_ppabspi1 = data['gstars_ppabspi1']
gstars_stud2 = data['gstars_stud2']
gstars_grade2 = data['gstars_grade2']
gstars_sec2 = data['gstars_sec2']
gstars_ppabspi2 = data['gstars_ppabspi2']
gstars_stud3 = data['gstars_stud3']
gstars_grade3= data['gstars_grade3']
gstars_sec3 = data['gstars_sec3']
gstars_ppabpsi3 = data['gstars_ppabspi3']

memoryIA = data['memoryIA']
visualpIA = data['visualpIA']
focusatIA = data['focusatIA']
probsolIA = data['probsolIA']
linguisIA = data['linguisIA']
memoryPPA = data['memoryPPA']
visualpPPA = data['visualpPPA']
focusatPPA = data['focusatPPA']
probsolPPA = data['probsolPPA']
linguisPPA = data['linguisPPA']

memIA20 = data['memIA20']
memIA40 = data['memIA40']
memIA60 = data['memIA60']
memIA80 = data['memIA80']
memIA100 = data['memIA100']
memPPA20 = data['memPPA20']
memPPA40 = data['memPPA40']
memPPA60 = data['memPPA60']
memPPA80 = data['memPPA80']
memPPA100 = data['memPPA100']

memgraIA1 = data['memgraIA1']
memgraIA2 = data['memgraIA2']
memgraIA3 = data['memgraIA3']
memgraIA4 = data['memgraIA4']
memgraIA5 = data['memgraIA5']
memgraPA1 = data['memgraPA1']
memgraPA2 = data['memgraPA2']
memgraPA3 = data['memgraPA3']
memgraPA4 = data['memgraPA4']
memgraPA5 = data['memgraPA5']

vispIA20 = data['vispIA20']
vispIA40 = data['vispIA40']
vispIA60 = data['vispIA60']
vispIA80 = data['vispIA80']
vispIA81 = data['vispIA81']
vispPA20 = data['vispPA20']
vispPA40 = data['vispPA40']
vispPA60 = data['vispPA60']
vispPA80 = data['vispPA80']
vispPA81 = data['vispPA81']

visgraIA1 = data['visgraIA1']
visgraIA2 = data['visgraIA2']
visgraIA3 = data['visgraIA3']
visgraIA4 = data['visgraIA4']
visgraIA5 = data['visgraIA5']
visgraPA1 = data['visgraPA1']
visgraPA2 = data['visgraPA2']
visgraPA3 = data['visgraPA3']
visgraPA4 = data['visgraPA4']
visgraPA5 = data['visgraPA5']

focuatIA_20 = data['focuatIA_20']
focuatIA_40 = data['focuatIA_40']
focuatIA_60 = data['focuatIA_60']
focuatIA_80 = data['focuatIA_80']
focuatIA_100 = data['focuatIA_100']
focuatPA_20 = data['focuatPA_20']
focuatPA_40 = data['focuatPA_40']
focuatPA_60 = data['focuatPA_60']
focuatPA_80 = data['focuatPA_80']
focuatPA_100 = data['focuatPA_100']

focgraIA1 = data['focgraIA1']
focgraIA2 = data['focgraIA2']
focgraIA3 = data['focgraIA3']
focgraIA4 = data['focgraIA4']
focgraIA5 = data['focgraIA5']
focgraPA1 = data['focgraPA1']
focgraPA2 = data['focgraPA2']
focgraPA3 = data['focgraPA3']
focgraPA4 = data['focgraPA4']
focgraPA5 = data['focgraPA5']

prosolIA_20 = data['prosolIA_20']
prosolIA_40 = data['prosolIA_40']
prosolIA_60 = data['prosolIA_60']
prosolIA_80 = data['prosolIA_80']
prosolIA_100 = data['prosolIA_100']
prosolPA_20 = data['prosolPA_20']
prosolPA_40 = data['prosolPA_40']
prosolPA_60 = data['prosolPA_60']
prosolPA_80 = data['prosolPA_80']
prosolPA_100 = data['prosolPA_100']

prograIA1 = data['prograIA1']
prograIA2 = data['prograIA2']
prograIA3 = data['prograIA3']
prograIA4 = data['prograIA4']
prograIA5 = data['prograIA5']
prograPA1 = data['prograPA1']
prograPA2 = data['prograPA2']
prograPA3 = data['prograPA3']
prograPA4 = data['prograPA4']
prograPA5 = data['prograPA5']

linguiIA_20 = data['linguiIA_20']
linguiIA_40 = data['linguiIA_40']
linguiIA_60 = data['linguiIA_60']
linguiIA_80 = data['linguiIA_80']
linguiIA_100 = data['linguiIA_100']
linguiPA_20 = data['linguiPA_20']
linguiPA_40 = data['linguiPA_40']
linguiPA_60 = data['linguiPA_60']
linguiPA_80 = data['linguiPA_80']
linguiPA_100 = data['linguiPA_100']

lingraIA1 = data['lingraIA1']
lingraIA2 = data['lingraIA2']
lingraIA3 = data['lingraIA3']
lingraIA4 = data['lingraIA4']
lingraIA5 = data['lingraIA5']
lingraPA1 = data['lingraPA1']
lingraPA2 = data['lingraPA2']
lingraPA3 = data['lingraPA3']
lingraPA4 = data['lingraPA4']
lingraPA5 = data['lingraPA5']


grade3avgscbel8 = data['grade3avgscbel8']
grade4avgscbel8 = data['grade4avgscbel8']
grade5avgscbel8 = data['grade5avgscbel8']
grade6avgscbel8 = data['grade6avgscbel8']
grade7avgscbel8 = data['grade7avgscbel8']
grade8avgscbel8 = data['grade8avgscbel8']
grade3avgsabove8 = data['grade3avgsabove8']
grade4avgsabove8 = data['grade4avgsabove8']
grade5avgsabove8 = data['grade5avgsabove8']
grade6avgsabove8 = data['grade6avgsabove8']
grade7avgsabove8 = data['grade7avgsabove8']
grade8avgsabove8 = data['grade8avgsabove8']

subskillgraph1 = data['subskillgraph1']
subskillgraph2 = data['subskillgraph2']
subskillgraph3 = data['subskillgraph3']
subskillgraph4 = data['subskillgraph4']
subskillgraph5 = data['subskillgraph5']
subskillgraph6 = data['subskillgraph6']
subskillgraph7 = data['subskillgraph7']
subskillgraph8 = data['subskillgraph8']
subskillgraph9 = data['subskillgraph9']
subskillgraph10 = data['subskillgraph10']
subskillgraph11 = data['subskillgraph11']
subskillgraph12 = data['subskillgraph12']
subskillgraph13 = data['subskillgraph13']
subskillgraph14 = data['subskillgraph14']
subskillgraph15 = data['subskillgraph15']
subskillgraph16 = data['subskillgraph16']
subskillgraph17 = data['subskillgraph17']
subskillgraph18 = data['subskillgraph18']
subskillgraph19 = data['subskillgraph19']
subskillgraph20 = data['subskillgraph20']
subskillgraph21 = data['subskillgraph21']

bspitop_studname1 = data['bspitop_studname1']
bspitop_grade1 = data['bspitop_grade1']
bspitop_sec1 = data['bspitop_sec1']
bspitop_bspi1 = data['bspitop_bspi1']
bspitop_studname2 = data['bspitop_studname2']
bspitop_grade2 = data['bspitop_grade2']
bspitop_sec2 = data['bspitop_sec2']
bspitop_bspi2 = data['bspitop_bspi2']
bspitop_studname3 = data['bspitop_studname3']
bspitop_grade3 = data['bspitop_grade3']
bspitop_sec3 = data['bspitop_sec3']
bspitop_bspi3 = data['bspitop_bspi3']

memtoppers_name1 = data['memtoppers_name1']
memtoppers_grade1 = data['memtoppers_grade1']
memtoppers_sec1 = data['memtoppers_sec1']
memtoppers_score1 = data['memtoppers_score1']
memtoppers_name2 = data['memtoppers_name2']
memtoppers_grade2 = data['memtoppers_grade2']
memtoppers_sec2 = data['memtoppers_sec2']
memtoppers_score2 = data['memtoppers_score2']
memtoppers_name3 = data['memtoppers_name3']
memtoppers_grade3 = data['memtoppers_grade3']
memtoppers_sec3 = data['memtoppers_sec3']
memtoppers_score3 = data['memtoppers_score3']


vptoppers_name1 = data['vptoppers_name1']
vptoppers_grade1 = data['vptoppers_grade1']
vptoppers_sec1 = data['vptoppers_sec1']
vptoppers_score1 = data['vptoppers_score1']
vptoppers_name2 = data['vptoppers_name2']
vptoppers_grade2 = data['vptoppers_grade2']
vptoppers_sec2 = data['vptoppers_sec2']
vptoppers_score2 = data['vptoppers_score2']
vptoppers_name3 = data['vptoppers_name3']
vptoppers_grade3 = data['vptoppers_grade3']
vptoppers_sec3 = data['vptoppers_sec3']
vptoppers_score3 = data['vptoppers_score3']

FAtoppers_name1 = data['FAtoppers_name1']
FAtoppers_grade1 = data['FAtoppers_grade1']
FAtoppers_sec1 = data['FAtoppers_sec1']
FAtoppers_score1 = data['FAtoppers_score1']
FAtoppers_name2 = data['FAtoppers_name2']
FAtoppers_grade2 = data['FAtoppers_grade2']
FAtoppers_sec2 = data['FAtoppers_sec2']
FAtoppers_score2 = data['FAtoppers_score2']
FAtoppers_name3 = data['FAtoppers_name3']
FAtoppers_grade3 = data['FAtoppers_grade3']
FAtoppers_sec3 = data['FAtoppers_sec3']
FAtoppers_score3 = data['FAtoppers_score3']

PSoltoppers_name1 = data['PSoltoppers_name1']
PSoltoppers_grade1 = data['PSoltoppers_grade1']
PSoltoppers_sec1 = data['PSoltoppers_sec1']
PSoltoppers_score1 = data['PSoltoppers_score1']
PSoltoppers_name2 = data['PSoltoppers_name2']
PSoltoppers_grade2 = data['PSoltoppers_grade2']
PSoltoppers_sec2 = data['PSoltoppers_sec2']
PSoltoppers_score2 = data['PSoltoppers_score2']
PSoltoppers_name3 = data['PSoltoppers_name3']
PSoltoppers_grade3 = data['PSoltoppers_grade3']
PSoltoppers_sec3 = data['PSoltoppers_sec3']
PSoltoppers_score3 = data['PSoltoppers_score3']

Lingtoppers_name1 = data['Lingtoppers_name1']
Lingtoppers_grade1 = data['Lingtoppers_grade1']
Lingtoppers_sec1 = data['Lingtoppers_sec1']
Lingtoppers_score1 = data['Lingtoppers_score1']
Lingtoppers_name2 = data['Lingtoppers_name2']
Lingtoppers_grade2 = data['Lingtoppers_grade2']
Lingtoppers_sec2 = data['Lingtoppers_sec2']
Lingtoppers_score2 = data['Lingtoppers_score2']
Lingtoppers_name3 = data['Lingtoppers_name3']
Lingtoppers_grade3 = data['Lingtoppers_grade3']
Lingtoppers_sec3 = data['Lingtoppers_sec3']
Lingtoppers_score3 = data['Lingtoppers_score3']

#First page insertion
from docx import Document
from docx.shared import Inches
from docx.shared import Cm
doc = Document('ppa_report.docx')
section = doc.sections[0]
header = section.header
for paragraph in header.paragraphs:
    for run in paragraph.runs:
        run.clear() 
section.different_first_page_header_footer = False
left_margin = section.left_margin
right_margin = section.right_margin
page_width = section.page_width - left_margin - right_margin
page_height = section.page_height
image_path = 'firstpage.png' 
doc.add_picture(image_path, width=page_width, height=page_height)
doc.save('ppa_report.docx')
print("First page added")

#Adding header
from docx import Document
from docx.shared import Inches
from docx.shared import Cm
doc = Document('ppa_report.docx')
section = doc.sections[0]
section.header_distance = Cm(0)  
left_margin = section.left_margin
right_margin = section.right_margin
width = section.page_width - left_margin - right_margin
header = section.header
paragraph = header.paragraphs[0]
run = paragraph.add_run()  
run.add_picture('header.png', width=width)  
doc.save('ppa_report.docx')
print("Header added")

#Text insertion and Center Alignment
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt  
doc = Document('ppa_report.docx')  
para = doc.add_paragraph(' Contents ')
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.runs[0]   
run.bold = True  
run.font.size = Pt(14)  
doc.save('ppa_report.docx')  

# Table creation 
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
data = [
    ["1", "The Impact of SKILLANGELS 2023-2024", "3"],
    ["1.1", "Aligning SKILLANGELS with other academic records", "4"],
    ["2", "Program Overview", "4"],
    ["2.1", "Skill Progression & Brain Skill Power Index (BSPI) comparison", "6"],
    ["2.2", "How they performed over the year", "7"],
    ["3", "A high-level overview of Assessment results", "8"],
    ["3.1", "BSPI comparison across grades", "8"],
    ["3.2", "BSPI improvement for each Grade", "9"],
    ["3.3", "Highly gifted students", "10"],
    ["3.4", "Future Achievers", "10"],
    ["3.5", "Growing Stars", "12"],
    ["4", "Skill Performance – Initial Vs. Post Program Assessment", "13"],
    ["5", "Do you know how well they did?", "19"],
    ["6", "Sub-Skill Score Analysis", "20"],
    ["7", "Achievers in PPA 2023-24", "21"],
    ["7.1", "BSPI Toppers by Grade & Section", "2.2"],
    ["7.2", "Skill Toppers by Grade & Section", "23"],
    ["8", "Final Inference", "26"],
    ["9", "Next level considerations", "26"],
    ["10", "Other Links", "27"]
]
doc = Document('ppa_report.docx')
table = doc.add_table(rows=21, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Sl No'
hdr_cells[1].text = 'Details'
hdr_cells[2].text = 'Page No'
for cell in hdr_cells:
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0]
    run.bold = True
    run.font.size = Pt(12)
for row in range(1, 21):  
    row_cells = table.rows[row].cells
    if row - 1 < len(data):  
        row_cells[0].text = data[row - 1][0]  
        row_cells[1].text = data[row - 1][1]  
        row_cells[2].text = data[row - 1][2]  
    else:
        row_cells[0].text = "N/A"
        row_cells[1].text = "N/A"
        row_cells[2].text = "N/A"
for row in table.rows:
    row.cells[0].width = Pt(70)
    row.cells[1].width = Pt(380)
    row.cells[2].width = Pt(70)
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(12)
def set_borders(table):
    for row in table.rows:
        for cell in row.cells:
            cell_xml = cell._element
            tc_pr = cell_xml.find(qn('w:tcPr'))
            if tc_pr is None:
                tc_pr = OxmlElement('w:tcPr')
                cell_xml.append(tc_pr)
            borders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')  
                border.set(qn('w:space'), '0')
                border.set(qn('w:sz'), '4')  # Border size
                borders.append(border)
            tc_pr.append(borders)
set_borders(table)
doc.save('ppa_report.docx')  
print("Table created")

# PAGE 2
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph('Annual Performance Report 2023-24')
para.alignment = WD_ALIGN_PARAGRAPH.CENTER  
run = para.runs[0]
run.bold = True
run.font.size = Pt(18)  
run.font.color.rgb = RGBColor(0, 112, 192)  
para = doc.add_paragraph('1. The Impact of SKILLANGELS')
run = para.runs[0]
run.font.size = Pt(14) 
run.font.color.rgb = RGBColor(45, 116, 181)

para = doc.add_paragraph('The academic year 2023-24 has been challenging in many respects, '
                          'with schools trying to find the right blend of skill-building and concept progression '
                          'at every stage and grade. EDSIX BRAINLAB has strengthened the efforts of ')
run = para.add_run(schoolname)
run.bold = True  
para.add_run(' by getting students on the SKILLANGELS platform to perform at their cognitive best. ')
para = doc.add_paragraph('Before the details of this impact are shared, we are happy to present the most relevant '
             'and essential findings of your students.')
for run in para.runs:
    run.font.size = Pt(12) 
doc.save('ppa_report.docx')

#LOGO INSERTION
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document('ppa_report.docx')  
para = doc.add_paragraph()
run = para.add_run()
logo_path = r'C:\Users\varsh\OneDrive\Sathyabama\Edsix\New folder\logo_image.png'  
run.add_picture(logo_path, width=Inches(2)) 
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save('ppa_report.docx')  
print("Logo inserted")

#Next para
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph('These critical data points highlight the impact that SKILLANGELS has had on the students at ')  
run = para.add_run(schoolname)
run.bold = True  
para.add_run('.')
run = para.add_run('This data has also highlighted the following:')
for run in para.runs:
    run.font.size = Pt(12) 
doc.save('ppa_report.docx')
print("Text, images are inserted successfully")

#Box Creation
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx
doc = Document('ppa_report.docx')
table = doc.add_table(rows=1, cols=1)
cell = table.cell(0, 0)
points = ["Performance is interrelated with continuous student-engagement.",
          "Students can climb the Cognitive Ladder effectively & steadily.",
          "Progression in multiple skillsets is evident and measurable."
]
for point in points:
    para = cell.add_paragraph(point, style='List Bullet')  
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.style.font.size = Pt(12)
cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Horizontal center
cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = docx.oxml.OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = docx.oxml.OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  # Solid line
            border.set(qn('w:sz'), '6')  
            border.set(qn('w:space'), '0')  
            border.set(qn('w:color'), 'FFA500')  
            borders.append(border)
        tc_pr.append(borders)
table.alignment = WD_ALIGN_PARAGRAPH.CENTER  
doc.save('ppa_report.docx')
print("Table created")

#Next Para
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph("Edsix team is proud to present this evidence-based report, which will pave way for further understanding and use of cognitive skilling that is essential to future-proof today's young learners.")
run = para.runs[0]
run.font.size = Pt(12)
doc.save('ppa_report.docx') 

#PAGE 3
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph('1.1 Aligning SKILLANGELS with other academic records:') 
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = para.runs[0]
run.font.size = Pt(14) 
run.font.color.rgb = RGBColor(45, 116, 181)
para = doc.add_paragraph('While highlighting the impact of this program, it would be relevant to note that Cognitive Upskilling gains a high level of attention and importance in the context of the current CBSE curriculum. The shift towards the Competency-based Learning & Assessment model that has been advocated by the')
run = para.add_run(' New Education Policy')
run.bold = True  
para.add_run('and the CBSE ensures that every child is given the support and the opportunity to enhance')
run = para.add_run(' Higher Order Thinking Skills')
run.bold = True
para.add_run(', thereby working towards competency and mastery in a subject. All teaching and learning efforts are outcome-based, leading to the need for a program that can be evidence for measurable achievements. With this program and quantifiable results, there is a clear indication of the continuum that every child moves on to help multiple stakeholders keep track at every stage.')
para = doc.add_paragraph("The SKILLANGELS CAUP scaffolds the learning process by providing an engaging and supportive platform in the learning cycle. For this precise reason, the data generated in this report should be viewed along with other academic records to get a comprehensive view of the student's achievements.")
para = doc.add_paragraph('Important terms in this report')
run = para.runs[0]
run.bold = True
for run in para.runs:
    run.font.size = Pt(12) 
doc.save('ppa_report.docx') 

#Table
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
doc = Document('ppa_report.docx')
table = doc.add_table(rows=3, cols=2)
table.cell(0, 0).text = 'CAUP'
table.cell(0, 1).text = 'COGNITIVE ABILITIES UPSKILLING PROGRAM'
table.cell(1, 0).text = 'IA & PPA'
table.cell(1, 1).text = 'Initial Assessment & Post Program Assessment'
table.cell(2, 0).text = 'BSPI'
table.cell(2, 1).text = 'BRAIN SKILL POWER INDEX- An average measure of Cognitive Brain Skills across 5 core areas: Memory, Visual Processing, Focus & Attention, Problem Solving, and Linguistics'
for row in table.rows:
    row.cells[0].paragraphs[0].runs[0].bold = True
for column in table.columns:
    for cell in column.cells:
        cell.width = Inches(4) 
for row in table.rows:
    row.cells[0].width = Pt(70)  
    row.cells[1].width = Pt(350)
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single') 
            border.set(qn('w:sz'), '4') 
            border.set(qn('w:space'), '0')  
            border.set(qn('w:color'), '000000')  #Black color
            borders.append(border)
        tc_pr.append(borders)
doc.save("ppa_report.docx") 
print("Table inserted")

#
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('2.Program Overview')
run = para.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(45, 116, 181)
doc.save('ppa_report.docx')

#Table Creation
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
doc = Document('ppa_report.docx')
table = doc.add_table(rows=6, cols=2)
table.cell(0, 0).text = 'Program name:'
table.cell(0, 1).text = 'SKILLANGELS – COGNITIVE ABILITIES UPSKILLING PROGRAM'
table.cell(1, 0).text = 'Program Start Date:'
table.cell(1, 1).text =  programstartdate
table.cell(2, 0).text = 'Number of students enrolled'
table.cell(2, 1).text =  numberofstudentsenrolled
table.cell(3, 0).text = 'Number of students who have completed the PPA'
table.cell(3, 1).text =  noofstudentsdoneppa
table.cell(4, 0).text = 'Grades in which the program has been conducted'
table.cell(4, 1).text =  gradesconducted
table.cell(5, 0).text = 'Reporting Period:'
table.cell(5, 1).text =  reportingperiod
for column in table.columns:
    for cell in column.cells:
        cell.width = Inches(4) 
for row in table.rows:
    row.cells[0].width = Pt(350)  
    row.cells[1].width = Pt(250)
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single') 
            border.set(qn('w:sz'), '4') 
            border.set(qn('w:space'), '0')  
            border.set(qn('w:color'), '000000')  #Black color
            borders.append(border)
        tc_pr.append(borders)
column_color = 'FFE499'  
for row in range(6):
    cell = table.cell(row, 0)  
    cell_xml = cell._element
    tc_pr = cell_xml.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), column_color)
    tc_pr.append(shading)
doc.save("ppa_report.docx") 
print("Table created")

#PAGE 4
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph('At EDSIX BRAINLAB we work to help children develop cognitive skills and become 21st century learners. We take pride in our collaboration with ')  
run = para.add_run(schoolname)
run.bold = True  
para.add_run(',')
run = para.add_run(' marking a significant endeavor. This aims to contribute to academic excellence and community engagement.')
para = doc.add_paragraph('')
run = para.add_run('Initial Assessment')
run.bold = True
para.add_run(' (IA) was conducted in June 2023 to establish the baseline of the student’s skill level at the start of the ')
run = para.add_run('SKILLANGELS - COGNITIVE ABILITIES UPSKILLING PROGRAM (CAUP). Post-Program Assessment')
run.bold = True
para.add_run('(PPA) was conducted at the end of the school year in March 2024. The progress of every learner is tracked right through the program and the progress can be observed on the Admin Portal.')
para = doc.add_paragraph('')
run = para.add_run('(PPA) was conducted at the end of the school year in March 2024. The progress of every learner is tracked right through the program and the progress can be observed on the Admin Portal.')
run.bold = True
for run in para.runs:
    run.font.size = Pt(12) 
doc.save('ppa_report.docx') 

#LOGO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document('ppa_report.docx')  
para = doc.add_paragraph()
run = para.add_run()
logo_path = r'C:\Users\varsh\OneDrive\Sathyabama\Edsix\New folder\logo2.png'  
run.add_picture(logo_path, width=Inches(7)) 
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save('ppa_report.docx')   
print("Logo inserted")

#Table
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
table = doc.add_table(rows=2, cols=3)
table.cell(0, 0).text = 'Minutes Trained'
table.cell(0, 1).text = 'Puzzles Attempted'
table.cell(0, 2).text = 'Puzzles Solved'
table.cell(1, 0).text =  minutestrained
table.cell(1, 1).text =  puzzlesattempted
table.cell(1, 2).text =  puzzlesolved
table.autofit = False  
col_width = Inches(2.0) 
for column in table.columns:
    for cell in column.cells:
        cell.width = col_width
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(14)
colors = ['D99594','FABC8F', 'C2D69B']  
for col_index, column in enumerate(table.columns):
    for cell in column.cells:
        cell_xml = cell._element
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), colors[col_index])  
        cell_xml.get_or_add_tcPr().append(shading) 
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  
            border.set(qn('w:sz'), '4') 
            border.set(qn('w:space'), '0')  
            border.set(qn('w:color'), '000000')  
            borders.append(border)
        tc_pr.append(borders)
doc.save("ppa_report.docx") 
print("Table inserted")

#Next topic (PAGE 4)
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph('2.1 Skill Progression & Brain Skill Power Index (BSPI) Comparison')
run = para.runs[0]
run.bold = True 
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(45, 116, 181)
doc.save('ppa_report.docx')

# Table
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
doc = Document('ppa_report.docx')
table = doc.add_table(rows=7, cols=4)
table.cell(0, 0).text = 'BSPI Score Range'
table.cell(0, 1).text = 'Description'
table.cell(0, 2).text = 'Number of students in IA'
table.cell(0, 3).text = 'Number of students in PPA'
table.cell(1, 0).text = '> 80'
table.cell(1, 1).text = 'Prodigy! This is fantastic. Keep it up and take it further!'
table.cell(1, 2).text =  bspiscorerangeIA_80
table.cell(1, 3).text =  bspiscorerangePPA_80
table.cell(2, 0).text = '> 60 and <= 80'
table.cell(2, 1).text = 'You have a spark. A little more determination and you will be there!'
table.cell(2, 2).text = bspiscorerangeIA_60 
table.cell(2, 3).text = bspiscorerangePPA_60
table.cell(3, 0).text = '> 40 and <= 60 '
table.cell(3, 1).text = 'The intervention will enhance academic performance. A slight push is all that is needed!'
table.cell(3, 2).text = bspiscorerangeIA_40
table.cell(3, 3).text = bspiscorerangePPA_40
table.cell(4, 0).text = ' > 20 and <= 40'
table.cell(4, 1).text = 'Needs Immediate Intervention. You can surely do it! Mistakes are the steppingstones to success.'
table.cell(4, 2).text = bspiscorerangeIA_20
table.cell(4, 3).text = bspiscorerangePPA_20
table.cell(5, 0).text = ' <= 20 '
table.cell(5, 1).text = "Rigorous and immediate intervention is needed. But don't feel low – You can certainly improve!"
table.cell(5, 2).text = bspiscorerangeIA_10
table.cell(5, 3).text = bspiscorerangePPA_10
table.cell(6, 0).text = ' '
table.cell(6, 1).text = 'Number of students who have completed the Program'
table.cell(6, 2).text = noofstudIA
table.cell(6, 3).text = noofstudIA
for column in table.columns:
    for cell in column.cells:
        cell.width = Inches(4) 
for row in [0, 6]:  
    for cell in table.rows[row].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)
for row in table.rows:
    row.cells[0].width = Pt(100)  
    row.cells[1].width = Pt(250)
    row.cells[2].width = Pt(100)
    row.cells[3].width = Pt(100)
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single') 
            border.set(qn('w:sz'), '4') 
            border.set(qn('w:space'), '0')  
            border.set(qn('w:color'), '000000')  #Black colour
            borders.append(border)
        tc_pr.append(borders)
row_colors = ['9ACC58', 'FDC850', 'FAA717', 'EB795F', 'EE4646']  
for i in range(1, 6):  
    for cell in table.rows[i].cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), row_colors[i-1]) 
        tc_pr.append(shading)
doc.save("ppa_report.docx") 
print("Table inserted")

# 
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph("The remarkable improvement observed in students' performance is evident. Initially, 114 students, representing 35% of the total, were classified in the top bands of BSPI during the initial assessment. However, this figure rose to 234 students, accounting for 71% of the total, in the post-assessment phase. This substantial increase suggests a notable enhancement in students' abilities over the course of the assessment period.")
run = para.runs[0]
run.font.size = Pt(12)
doc.save('ppa_report.docx') 

#
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('We congratulate the students and teachers on this excellent performance.')
para.alignment = WD_ALIGN_PARAGRAPH.CENTER  
run = para.runs[0]
run.bold = True  
run.font.size = Pt(14)  
run.font.color.rgb = RGBColor(0, 112, 192)  
doc.save('ppa_report.docx') 

#logo insertion
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document('ppa_report.docx')  
para = doc.add_paragraph()
run = para.add_run()
logo_path = r'C:\Users\varsh\OneDrive\Sathyabama\Edsix\New folder\logo_image.png'  
run.add_picture(logo_path, width=Inches(1)) 
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save('ppa_report.docx')   

#Next page
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph('2.2 How they performed over the year.')
run = para.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(45, 116, 181)
doc.save('ppa_report.docx')

#
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('The overall school average BSPI has increased from ')  
run = para.add_run( schoolavgbspi)
run.bold = True  
run.font.color.rgb = RGBColor(0, 175, 80)
para.add_run('.')
para = doc.add_paragraph('Continuous tracking of student engagement and performance has shown that every student who consistently uses the platform is able to move ahead steadily and achieve significantly higher scores from the beginning to the end of the program.')
para = doc.add_paragraph("In expressing gratitude for the teachers' commendable efforts in this academic year, we encourage sustained and uplifting support for every student to actively participate and engage with the program. This would definitively show positive results not just in cognitive skill building but in academic performance as well.")
for run in para.runs:
    run.font.size = Pt(12) 
doc.save('ppa_report.docx') 
print("Text inserted")

#logo
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document('ppa_report.docx')  
para = doc.add_paragraph()
run = para.add_run()
logo_path = r'C:\Users\varsh\OneDrive\Sathyabama\Edsix\New folder\logo_image.png'  
run.add_picture(logo_path, width=Inches(2)) 
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save('ppa_report.docx')   

#topic
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph('3 A high-level overview of Assessment results  3.1 BSPI comparison across grades: ')
run = para.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(45, 116, 181)
doc.save('ppa_report.docx')

# next para
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('Since the ')  
run = para.add_run('IA')
run.bold = True  
para.add_run(',students have improved in their overall skill proficiency on average, which is indicated in the table below. Students who have completed fewer than eight sessions were not considered for this analysis since they would not have had enough practice to make a meaningful comparison.')
para = doc.add_paragraph('Over the course of the academic year, noticeable improvement in the average BSPI has been evident across ALL grades from the beginning to the end. ')
for run in para.runs:
    run.font.size = Pt(12) 
doc.save('ppa_report.docx') 

#GRAPH INSERTION
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
import io
def rgb_to_normalized(rgb):
    return tuple([x / 255.0 for x in rgb])
background_color = rgb_to_normalized((32, 56, 100))  
bar1_color = rgb_to_normalized((154, 87, 205))  
bar2_color = rgb_to_normalized((146, 208, 80)) 
items = ['III', 'IV', 'V', 'VI', 'VII', 'VIII']
initial_scores = [bspigrowthgraphIA1, bspigrowthgraphIA2, bspigrowthgraphIA3, bspigrowthgraphIA4, bspigrowthgraphIA5, bspigrowthgraphIA6]  
post_scores = [bspigrowthgraphPPA1, bspigrowthgraphPPA2, bspigrowthgraphPPA3, bspigrowthgraphPPA4, bspigrowthgraphPPA5, bspigrowthgraphPPA6]   
x = np.arange(len(items))  
width = 0.25  
fig, ax = plt.subplots(figsize=(9, 5))
plt.gcf().set_facecolor(background_color)  
ax.set_facecolor(background_color) 
gap = 0.02
bar1 = ax.bar(x - width/2 - gap, initial_scores, width, label='Initial Assessment', color=bar1_color)  
bar2 = ax.bar(x + width/2 + gap, post_scores, width, label='Post Program Assessment', color=bar2_color) 
ax.set_title('BSPI growth-Initial vs Post Program Assessment', color='white')  
ax.set_xticks(x)
ax.set_xticklabels(items, color='white') 
ax.tick_params(axis='y', colors='white')  
ax.tick_params(axis='x', colors='white')  
ax.legend(facecolor='gray', edgecolor='white', loc='upper right', fontsize=10, framealpha=0.7)  
for i, rect in enumerate(bar1):
    height = rect.get_height()
    ax.text(rect.get_x() + rect.get_width() / 2, height, str(initial_scores[i]), ha='center', va='bottom', color='white')
for i, rect in enumerate(bar2):
    height = rect.get_height()
    ax.text(rect.get_x() + rect.get_width() / 2, height, str(post_scores[i]), ha='center', va='bottom', color='white')
ax.grid(True, which='both', axis='y' , linestyle='-', color='black', linewidth=0.6, alpha=0.2)
plt.tight_layout()
img_stream = io.BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0) 
doc = Document('ppa_report.docx')
table = doc.add_table(rows=1, cols=1)
cell = table.cell(0, 0)
cell.paragraphs[0].clear()  
cell.paragraphs[0].add_run().add_picture(img_stream, width=Inches(6)) 
cell.alignment = 1  
doc.save('ppa_report.docx')
img_stream.close() 
print("Graph inserted")

#next para
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph("The visible improvement from initial assessment to post program assessment indicates a positive impact of the CAUP training, demonstrating its effectiveness in enhancing students' skill sets.")
para = doc.add_paragraph('This positive trend underscores the collective progress and proficiency development observed across various aspects of learning. Congratulations!')
for run in para.runs:
    run.font.size = Pt(12) 
doc.save('ppa_report.docx')

# topic
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph('3.2 BSPI improvement for each grade:')
run = para.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(45, 116, 181)
doc.save('ppa_report.docx')

#Logo insertion 
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
run = para.add_run()
logo_path = r'C:\Users\varsh\OneDrive\Sathyabama\Edsix\New folder\logo3.png'  
run.add_picture(logo_path, width=Inches(6)) 
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save('ppa_report.docx')   
print("Logo inserted")

#Next para
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('The data below shows the BSPI improvement for each grade.')
para = doc.add_paragraph('This report is analyzed based on the number of students who have taken both the ')  
run = para.add_run('IA')
run.bold = True
para.add_run(' as well as the ')
run = para.add_run('PPA.')
run.bold = True  
para = doc.add_paragraph('The total number of students who have taken both is ')
run = para.add_run( noofstudIA)
run.bold = True
para.add_run('.(as on')
run = para.add_run( reportingperiod)
run.bold = True
para.add_run(').')
for run in para.runs:
    run.font.size = Pt(12) 
doc.save('ppa_report.docx') 

#Table
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
doc = Document('ppa_report.docx')
table = doc.add_table(rows=3, cols=6)
table.cell(0, 0).text = 'Grade III '
table.cell(0, 1).text = '<= 20'
table.cell(0, 2).text = '21-40 '
table.cell(0, 3).text = '41-60 '
table.cell(0, 4).text = '61-80 '
table.cell(0, 5).text = '> 80 '
table.cell(1, 0).text = 'Initial Assessment '
table.cell(1, 1).text =  grade3IA20
table.cell(1, 2).text =  grade3IA40
table.cell(1, 3).text =  grade3IA60
table.cell(1, 4).text =  grade3IA80
table.cell(1, 5).text =  grade3IA100
table.cell(2, 0).text = 'Post Program Assessment'
table.cell(2, 1).text =  grade3PPA20
table.cell(2, 2).text =  grade3PPA40
table.cell(2, 3).text =  grade3PPA60
table.cell(2, 4).text =  grade3PPA80
table.cell(2, 5).text =  grade3PPA100
for row in [0, 2]:  
    cell = table.cell(row, 0) 
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
for row in table.rows:
    row.cells[0].width = Pt(150)  
    row.cells[1].width = Pt(60)  
    row.cells[2].width = Pt(60)
    row.cells[3].width = Pt(60)
    row.cells[4].width = Pt(60)
    row.cells[5].width = Pt(60)
background_colors = [
    RGBColor(238, 70, 70),   # Red
    RGBColor(235, 121, 96),  # Orange
    RGBColor(250, 172, 23),  # Yellow
    RGBColor(253, 194, 80),  # Light Yellow
    RGBColor(154, 204, 88)   # Green
]

def rgb_to_hex(rgb_color):
    return f'{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}'
for col_idx in range(1, 6):  
    cell = table.cell(0, col_idx)
    cell_text = cell.text
    cell.text = ""
    run = cell.paragraphs[0].add_run(cell_text)
    hex_color = rgb_to_hex(background_colors[col_idx - 1])
    cell._element.get_or_add_tcPr().append(
        OxmlElement('w:shd', {qn('w:fill'): hex_color})
    )
def set_cell_borders(cell):
    cell_xml = cell._element
    tc_pr = cell_xml.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  # Single border
        border.set(qn('w:sz'), '4')  # Border width
        border.set(qn('w:space'), '0')  # Space between borders
        border.set(qn('w:color'), '000000')  # Black color for borders
        borders.append(border)
    tc_pr.append(borders)
for row in table.rows:
    for cell in row.cells:
        set_cell_borders(cell)
doc.save('ppa_report.docx') 
print("Table inserted successfully")

#Table2
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
table = doc.add_table(rows=3, cols=6)
table.cell(0, 0).text = 'Grade IV '
table.cell(0, 1).text = '<= 20'
table.cell(0, 2).text = '21-40 '
table.cell(0, 3).text = '41-60 '
table.cell(0, 4).text = '61-80 '
table.cell(0, 5).text = '> 80 '
table.cell(1, 0).text = 'Initial Assessment '
table.cell(1, 1).text = grade4IA20
table.cell(1, 2).text = grade4IA40
table.cell(1, 3).text = grade4IA60
table.cell(1, 4).text = grade4IA80
table.cell(1, 5).text = grade4IA100
table.cell(2, 0).text = 'Post Program Assessment'
table.cell(2, 1).text = grade4PPA20
table.cell(2, 2).text = grade4PPA40
table.cell(2, 3).text = grade4PPA60
table.cell(2, 4).text = grade4PPA80
table.cell(2, 5).text = grade4PPA100
for row in [0, 2]:  
    cell = table.cell(row, 0)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
for row in table.rows:
    row.cells[0].width = Pt(150)  
    row.cells[1].width = Pt(60)  
    row.cells[2].width = Pt(60)
    row.cells[3].width = Pt(60)
    row.cells[4].width = Pt(60)
    row.cells[5].width = Pt(60)
for col_idx in range(6):  
    cell = table.cell(2, col_idx)  
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
background_colors = [
    RGBColor(238, 70, 70),   # Red
    RGBColor(235, 121, 96),  # Orange
    RGBColor(250, 172, 23),  # Yellow
    RGBColor(253, 194, 80),  # Light Yellow
    RGBColor(154, 204, 88)   # Green
]
def rgb_to_hex(rgb_color):
    return f'{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}'
for col_idx in range(1, 6):  
    cell = table.cell(0, col_idx)
    cell_text = cell.text
    cell.text = ""
    run = cell.paragraphs[0].add_run(cell_text)
    hex_color = rgb_to_hex(background_colors[col_idx - 1])
    cell._element.get_or_add_tcPr().append(
        OxmlElement('w:shd', {qn('w:fill'): hex_color})
    )
def set_cell_borders(cell):
    cell_xml = cell._element
    tc_pr = cell_xml.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  
        border.set(qn('w:sz'), '4')  
        border.set(qn('w:space'), '0')  
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tc_pr.append(borders)
for row in table.rows:
    for cell in row.cells:
        set_cell_borders(cell)
doc.save('ppa_report.docx') 
print("Table inserted")

#Table3
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
table = doc.add_table(rows=3, cols=6)
table.cell(0, 0).text = 'Grade V'
table.cell(0, 1).text = '<= 20'
table.cell(0, 2).text = '21-40 '
table.cell(0, 3).text = '41-60 '
table.cell(0, 4).text = '61-80 '
table.cell(0, 5).text = '> 80 '
table.cell(1, 0).text = 'Initial Assessment'
table.cell(1, 1).text = grade5IA20
table.cell(1, 2).text = grade5IA40
table.cell(1, 3).text = grade5IA60
table.cell(1, 4).text = grade5IA80
table.cell(1, 5).text = grade5IA100
table.cell(2, 0).text = 'Post Program Assessment'
table.cell(2, 1).text = grade5PPA20
table.cell(2, 2).text = grade5PPA40
table.cell(2, 3).text = grade5PPA60
table.cell(2, 4).text = grade5PPA80
table.cell(2, 5).text = grade5PPA100
for row in [0, 2]:  
    cell = table.cell(row, 0)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
for row in table.rows:
    row.cells[0].width = Pt(150)  
    row.cells[1].width = Pt(60)  
    row.cells[2].width = Pt(60)
    row.cells[3].width = Pt(60)
    row.cells[4].width = Pt(60)
    row.cells[5].width = Pt(60)
for col_idx in range(6): 
    cell = table.cell(2, col_idx)  
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
background_colors = [
    RGBColor(238, 70, 70),   # Red
    RGBColor(235, 121, 96),  # Orange
    RGBColor(250, 172, 23),  # Yellow
    RGBColor(253, 194, 80),  # Light Yellow
    RGBColor(154, 204, 88)   # Green
]
def rgb_to_hex(rgb_color):
    return f'{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}'
for col_idx in range(1, 6):  
    cell = table.cell(0, col_idx)
    cell_text = cell.text
    cell.text = ""
    run = cell.paragraphs[0].add_run(cell_text)
    hex_color = rgb_to_hex(background_colors[col_idx - 1])
    cell._element.get_or_add_tcPr().append(
        OxmlElement('w:shd', {qn('w:fill'): hex_color})
    )
def set_cell_borders(cell):
    cell_xml = cell._element
    tc_pr = cell_xml.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  
        border.set(qn('w:sz'), '4') 
        border.set(qn('w:space'), '0')  
        border.set(qn('w:color'), '000000') 
        borders.append(border)
    tc_pr.append(borders)
for row in table.rows:
    for cell in row.cells:
        set_cell_borders(cell)
doc.save('ppa_report.docx') 
print("Table inserted")

#Table 4
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
table = doc.add_table(rows=3, cols=6)
table.cell(0, 0).text = 'Grade VI'
table.cell(0, 1).text = '<= 20'
table.cell(0, 2).text = '21-40'
table.cell(0, 3).text = '41-60'
table.cell(0, 4).text = '61-80'
table.cell(0, 5).text = '> 80'
table.cell(1, 0).text = 'Initial Assessment'
table.cell(1, 1).text = grade6IA20
table.cell(1, 2).text = grade6IA40
table.cell(1, 3).text = grade6IA60
table.cell(1, 4).text = grade6IA80
table.cell(1, 5).text = grade6IA100
table.cell(2, 0).text = 'Post Program Assessment'
table.cell(2, 1).text = grade6PPA20
table.cell(2, 2).text = grade6PPA40
table.cell(2, 3).text = grade6PPA60
table.cell(2, 4).text = grade6PPA80
table.cell(2, 5).text = grade6PPA100
for row in [0, 2]:  
    cell = table.cell(row, 0)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
for row in table.rows:
    row.cells[0].width = Pt(150) 
    row.cells[1].width = Pt(60)  
    row.cells[2].width = Pt(60)
    row.cells[3].width = Pt(60)
    row.cells[4].width = Pt(60)
    row.cells[5].width = Pt(60)
for col_idx in range(6): 
    cell = table.cell(2, col_idx)  
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
background_colors = [
    RGBColor(238, 70, 70),   # Red
    RGBColor(235, 121, 96),  # Orange
    RGBColor(250, 172, 23),  # Yellow
    RGBColor(253, 194, 80),  # Light Yellow
    RGBColor(154, 204, 88)   # Green
]
def rgb_to_hex(rgb_color):
    return f'{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}'
for col_idx in range(1, 6):  
    cell = table.cell(0, col_idx)
    cell_text = cell.text
    cell.text = ""
    run = cell.paragraphs[0].add_run(cell_text)
    hex_color = rgb_to_hex(background_colors[col_idx - 1])
    cell._element.get_or_add_tcPr().append(
        OxmlElement('w:shd', {qn('w:fill'): hex_color})
    )
def set_cell_borders(cell):
    cell_xml = cell._element
    tc_pr = cell_xml.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  
        border.set(qn('w:sz'), '4')  
        border.set(qn('w:space'), '0')  
        border.set(qn('w:color'), '000000')  
        borders.append(border)
    tc_pr.append(borders)
for row in table.rows:
    for cell in row.cells:
        set_cell_borders(cell)
doc.save('ppa_report.docx') 
print("Table created")

#
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
table = doc.add_table(rows=3, cols=6)
table.cell(0, 0).text = 'Grade VII'
table.cell(0, 1).text = '<= 20'
table.cell(0, 2).text = '21-40 '
table.cell(0, 3).text = '41-60 '
table.cell(0, 4).text = '61-80 '
table.cell(0, 5).text = '> 80 '
table.cell(1, 0).text = 'Initial Assessment'
table.cell(1, 1).text = grade7IA20
table.cell(1, 2).text = grade7IA40
table.cell(1, 3).text = grade7IA60
table.cell(1, 4).text = grade7IA80
table.cell(1, 5).text = grade7IA100
table.cell(2, 0).text = 'Post Program Assessment'
table.cell(2, 1).text = grade7PPA20
table.cell(2, 2).text = grade7PPA40
table.cell(2, 3).text = grade7PPA60
table.cell(2, 4).text = grade7PPA80
table.cell(2, 5).text = grade7PPA100
for row in [0, 2]:  
    cell = table.cell(row, 0)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
for row in table.rows:
    row.cells[0].width = Pt(150)  
    row.cells[1].width = Pt(60)  
    row.cells[2].width = Pt(60)
    row.cells[3].width = Pt(60)
    row.cells[4].width = Pt(60)
    row.cells[5].width = Pt(60)
for col_idx in range(6): 
    cell = table.cell(2, col_idx)  
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
background_colors = [
    RGBColor(238, 70, 70),   # Red
    RGBColor(235, 121, 96),  # Orange
    RGBColor(250, 172, 23),  # Yellow
    RGBColor(253, 194, 80),  # Light Yellow
    RGBColor(154, 204, 88)   # Green
]

def rgb_to_hex(rgb_color):
    return f'{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}'
for col_idx in range(1, 6):  
    cell = table.cell(0, col_idx)
    cell_text = cell.text
    cell.text = ""
    run = cell.paragraphs[0].add_run(cell_text)
    hex_color = rgb_to_hex(background_colors[col_idx - 1])
    cell._element.get_or_add_tcPr().append(
        OxmlElement('w:shd', {qn('w:fill'): hex_color})
    )
def set_cell_borders(cell):
    cell_xml = cell._element
    tc_pr = cell_xml.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  
        border.set(qn('w:sz'), '4')  
        border.set(qn('w:space'), '0')  
        border.set(qn('w:color'), '000000')  
        borders.append(border)
    tc_pr.append(borders)
for row in table.rows:
    for cell in row.cells:
        set_cell_borders(cell)
doc.save('ppa_report.docx') 
print("Table inserted")

#
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
table = doc.add_table(rows=3, cols=6)
table.cell(0, 0).text = 'Grade VIII '
table.cell(0, 1).text = '<= 20'
table.cell(0, 2).text = '21-40 '
table.cell(0, 3).text = '41-60 '
table.cell(0, 4).text = '61-80 '
table.cell(0, 5).text = '> 80 '
table.cell(1, 0).text = 'Initial Assessment '
table.cell(1, 1).text = grade8IA20
table.cell(1, 2).text = grade8IA40
table.cell(1, 3).text = grade8IA60
table.cell(1, 4).text = grade8IA80
table.cell(1, 5).text = grade8IA100
table.cell(2, 0).text = 'Post Program Assessment'
table.cell(2, 1).text = grade8PPA20
table.cell(2, 2).text = grade8PPA40
table.cell(2, 3).text = grade8PPA60
table.cell(2, 4).text = grade8PPA80
table.cell(2, 5).text = grade8PPA100
for row in [0, 2]:  
    cell = table.cell(row, 0)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
for row in table.rows:
    row.cells[0].width = Pt(150)  # First column narrow
    row.cells[1].width = Pt(60)  # Middle column wide
    row.cells[2].width = Pt(60)
    row.cells[3].width = Pt(60)
    row.cells[4].width = Pt(60)
    row.cells[5].width = Pt(60)
for col_idx in range(6): 
    cell = table.cell(2, col_idx)  
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
background_colors = [
    RGBColor(238, 70, 70),   # Red
    RGBColor(235, 121, 96),  # Orange
    RGBColor(250, 172, 23),  # Yellow
    RGBColor(253, 194, 80),  # Light Yellow
    RGBColor(154, 204, 88)   # Green
]

def rgb_to_hex(rgb_color):
    return f'{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}'
for col_idx in range(1, 6):  
    cell = table.cell(0, col_idx)
    cell_text = cell.text
    cell.text = ""
    run = cell.paragraphs[0].add_run(cell_text)
    hex_color = rgb_to_hex(background_colors[col_idx - 1])
    cell._element.get_or_add_tcPr().append(
        OxmlElement('w:shd', {qn('w:fill'): hex_color})
    )
def set_cell_borders(cell):
    cell_xml = cell._element
    tc_pr = cell_xml.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  
        border.set(qn('w:sz'), '4')  
        border.set(qn('w:space'), '0')  
        border.set(qn('w:color'), '000000')  
        borders.append(border)
    tc_pr.append(borders)
for row in table.rows:
    for cell in row.cells:
        set_cell_borders(cell)
doc.save('ppa_report.docx') 
print("Table inserted")

#Topic
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph('3.3 Highly gifted students - Prodigies - > 80 PPA BSPI: ')
run = para.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(45, 116, 181)
doc.save('ppa_report.docx')

#Next para
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('Congratulations to our highly gifted student, the prodigies with PPA BSPI scores exceeding 80, on her promising improvement! Keep up the exceptional work and continue to shine brightly!')  
run = para.runs[0]
run.font.size = Pt(12)
doc.save('ppa_report.docx') 

# TABLE 
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
doc = Document('ppa_report.docx')
table = doc.add_table(rows=3, cols=5)
cell = table.cell(0, 2)
cell.text = '                                 Highly gifted students - Prodigies - > 80 PPA BSPI '
run.bold = True
table.cell(0, 0).merge(table.cell(0, 4))
cell = table.cell(0, 0)
cell.paragraphs[0].alignment = 1  
for run in cell.paragraphs[0].runs:
    run.bold = True
    run.font.size = Pt(14) 
    run.font.color.rgb = RGBColor(0, 0, 0) 
table.cell(1, 0).text = 'S.NO'
table.cell(1, 1).text = 'Student Name'
table.cell(1, 2).text = 'Grade'
table.cell(1, 3).text = 'Section'
table.cell(1, 4).text = 'PPA - BSPI'
table.cell(2, 0).text = '1'
table.cell(2, 1).text = highly_prodigiesname
table.cell(2, 2).text = highly_prodigiesgrade
table.cell(2, 3).text = highly_prodigiessec
table.cell(2, 4).text = highly_prodigiesBSPI
for row in [0, 1]:  
    for cell in table.rows[row].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
for row in table.rows:
    row.cells[0].width = Pt(30)  # First column narrow
    row.cells[1].width = Pt(250)  # Middle column wide
    row.cells[2].width = Pt(60)
    row.cells[3].width = Pt(60)
    row.cells[4].width = Pt(300)
background_color = RGBColor(146, 208, 80)
def rgb_to_hex(rgb_color):
    return f'{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}' 
hex_color = rgb_to_hex((146, 208, 80))  
for col_idx in range(5):  
    cell = table.cell(0, col_idx)
    cell_text = cell.text
    cell.text = ""
    run = cell.paragraphs[0].add_run(cell_text)
    cell._element.get_or_add_tcPr().append(
        OxmlElement('w:shd', {qn('w:fill'): hex_color})
    )
def set_cell_borders(cell, border_width=4):
    cell_xml = cell._element
    tc_pr = cell_xml.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  
        border.set(qn('w:sz'), str(border_width))  
        border.set(qn('w:space'), '0')  
        border.set(qn('w:color'), '000000')  
        borders.append(border)
    tc_pr.append(borders)
border_width = 6  
for row in table.rows:
    for cell in row.cells:
        set_cell_borders(cell, border_width)
doc.save('ppa_report.docx') 
print("Table inserted")

# Topic
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph('3.4  Future Achievers - 70-80 PPA BSPI:')
run = para.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(45, 116, 181)
doc.save('ppa_report.docx')

#Next para
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('Best wishes to our future achievers, those with PPA BSPI scores ranging from 70 to 80, for their significant improvement! With further training, they are destined to become the stars of tomorrow. Keep up the great work!')  
run = para.runs[0]
run.font.size = Pt(12)
doc.save('ppa_report.docx') 

#Table
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
doc = Document('ppa_report.docx')
table = doc.add_table(rows=24, cols=5)
cell = table.cell(0, 0)
cell.text = '                                FUTURE ACHIEVERS - 70 to 80 PPA BSPI  '
table.cell(0, 0).merge(table.cell(0, 4))
cell = table.cell(0, 0)
cell.paragraphs[0].alignment = 1 
for run in cell.paragraphs[0].runs:
    run.bold = True
    run.font.size = Pt(16) 
    run.font.color.rgb = RGBColor(0, 0, 0) 
table.cell(1, 0).text = 'S.NO'
table.cell(1, 1).text = 'Student Name'
table.cell(1, 2).text = 'Grade'
table.cell(1, 3).text = 'Section'
table.cell(1, 4).text = 'PPA - BSPI'
data = [
    ('1', FAstud1, FAgrade1, FAsec1, FAppa_bspi1),
    ('2', FAstud2, FAgrade2, FAsec2, FAppa_bspi2),
    ('4', FAstud3, FAgrade3, FAsec3, FAppa_bspi3),
    ('4', 'SHANMUGAPRIYA', '3', 'B', '77.6'),
    ('5', 'TEJASHWIN', '3', 'B', '72.2'),
    ('6', 'THANUSHREE', '3', 'B', '71'),
    ('7', 'SAKTHI KARNIKA', '4', 'A', '73.4'),
    ('8', 'HEMADHARSHINI', '4', 'B', '75.4'),
    ('9', 'DHARNEESH', '4', 'B', '72.8'),
    ('10', 'KESAV', '5', 'A', '79.2'),
    ('11', 'VISHNUVARTHAN', '5', 'A', '70.6'),
    ('12', 'THISANTH', '5', 'B', '75.8'),
    ('13', 'KRISHNA KANNIKA', '5', 'B', '74'),
    ('14', 'AKSHAYA KEERTHI', '6', 'A', '76.2'),
    ('15', 'VIDHARSHANA', '6', 'B', '79.2'),
    ('16', 'JOSHNI SREE', '6', 'B', '76.4'),
    ('17', 'NIVEDHA', '6', 'B', '72'),
    ('18', 'UMMUL AYISHA', '7', 'A', '79'),
    ('19', 'DHAANYA', '7', 'A', '75.6'),
    ('20', 'DEVESH KUMAR', '7', 'A', '73.8'),
    ('21', 'SANJAI', '7', 'A', '73.6'),
    ('22', 'GOPIKA', '7', 'A', '71.6')
]
for i, (sno, name, grade, section, ppa) in enumerate(data, start=2):
    table.cell(i, 0).text = sno
    table.cell(i, 1).text = name
    table.cell(i, 2).text = grade
    table.cell(i, 3).text = section
    table.cell(i, 4).text = ppa
for row in [0, 1]:  
    for cell in table.rows[row].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
for row in table.rows:
    row.cells[0].width = Pt(30)  # First column narrow
    row.cells[1].width = Pt(250)  # Middle column wide
    row.cells[2].width = Pt(60)
    row.cells[3].width = Pt(60)
    row.cells[4].width = Pt(300)
background_color = RGBColor(255, 192, 0)
def rgb_to_hex(rgb_color):
    return f'{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}'
hex_color = rgb_to_hex((255, 192, 0))  
for col_idx in range(5):  
    cell = table.cell(0, col_idx)
    cell_text = cell.text
    cell.text = ""
    run = cell.paragraphs[0].add_run(cell_text)
    cell._element.get_or_add_tcPr().append(
        OxmlElement('w:shd', {qn('w:fill'): hex_color})
    )
def set_cell_borders(cell, border_width=4):
    cell_xml = cell._element
    tc_pr = cell_xml.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_width))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tc_pr.append(borders)
border_width = 6
for row in table.rows:
    for cell in row.cells:
        set_cell_borders(cell, border_width)
doc.save('ppa_report.docx') 
print("Table inserted")


# Topic
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph('3.5 Growing Stars - <= 20 PPA BSPI:')
run = para.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(45, 116, 181)
doc.save('ppa_report.docx')

#Next para
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('These students (<= 20 BSPI) require immediate attention for the improvement. With proper training, they are certain to advance to higher bands like their peers. Wishing them the very best as they embark on their path to success!')  
run = para.runs[0]
run.font.size = Pt(12)
doc.save('ppa_report.docx') 

#Table
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
table = doc.add_table(rows=16, cols=5)
cell = table.cell(0, 0)
cell.text = '                               GROWING STARS- <= 20 PPA BSPI   '
table.cell(0, 0).merge(table.cell(0, 4))
cell = table.cell(0, 0)
cell.paragraphs[0].alignment = 1
for run in cell.paragraphs[0].runs:
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
table.cell(1, 0).text = 'S.NO'
table.cell(1, 1).text = 'Student Name'
table.cell(1, 2).text = 'Grade'
table.cell(1, 3).text = 'Section'
table.cell(1, 4).text = 'PPA - BSPI'
data = [
    ('1', gstars_stud1, gstars_grade1, gstars_sec1, gstars_ppabspi1),
    ('2', gstars_stud2, gstars_grade2, gstars_sec2, gstars_ppabspi2),
    ('3', gstars_stud3, gstars_grade3, gstars_sec3, gstars_ppabpsi3),
    ('4', 'RAMAVARSHAN', '5', 'A', '19.6'),
    ('5', 'SABARIVASAN', '6', 'B', '14.4'),
    ('6', 'SASHVINTH', '7', 'A', '16.4'),
    ('7', 'MLLER', '8', 'A', '18'),
    ('8', 'SUBASH', '8', 'A', '18.2'),
    ('9', 'PRAJITH', '8', 'A', '18.6'),
    ('10', 'SARVESH VIJAY', '8', 'B', '8'),
    ('11', 'RITHIKA SRI', '8', 'B', '17'),
    ('12', 'PARTHIBAN', '8', 'B', '17.2'),
    ('13', 'KAVIYA BHARATHI', '8', 'B', '17.8'),
    ('14', 'KABIL', '8', 'B', '19.8')
]
for i, (sno, name, grade, section, ppa) in enumerate(data, start=2):
    table.cell(i, 0).text = sno
    table.cell(i, 1).text = name
    table.cell(i, 2).text = grade
    table.cell(i, 3).text = section
    table.cell(i, 4).text = ppa
for row in [0, 1]:  
    for cell in table.rows[row].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
for row in table.rows:
    row.cells[0].width = Pt(30)  
    row.cells[1].width = Pt(250)  
    row.cells[2].width = Pt(60)
    row.cells[3].width = Pt(60)
    row.cells[4].width = Pt(300)
background_color = RGBColor(255, 192, 0)
def rgb_to_hex(rgb_color):
    return f'{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}'
hex_color = rgb_to_hex((255, 192, 0))  
for col_idx in range(5):  
    cell = table.cell(0, col_idx)
    cell_text = cell.text
    cell.text = ""
    run = cell.paragraphs[0].add_run(cell_text)
    cell._element.get_or_add_tcPr().append(
        OxmlElement('w:shd', {qn('w:fill'): hex_color})
    )
def set_cell_borders(cell, border_width=4):
    cell_xml = cell._element
    tc_pr = cell_xml.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_width))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tc_pr.append(borders)
border_width = 6
for row in table.rows:
    for cell in row.cells:
        set_cell_borders(cell, border_width)
doc.save('ppa_report.docx') 
print("text inserted")

# Topic
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph('4 Skill Performance – Initial Vs.Post Program Assessment  ')
run = para.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(45, 116, 181)
doc.save('ppa_report.docx')

#Logo
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document('ppa_report.docx')  
para = doc.add_paragraph()
run = para.add_run()
logo_path = r'C:\Users\varsh\OneDrive\Sathyabama\Edsix\New folder\logo4.png'  
run.add_picture(logo_path, width=Inches(6)) 
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save('ppa_report.docx')   
print("Logo inserted")

# Text center
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('The 5 essential skills that the Program tracks and measures.')
para.alignment = WD_ALIGN_PARAGRAPH.CENTER  
run = para.runs[0]
run.bold = True  
run.font.size = Pt(12)  
run.font.color.rgb = RGBColor(255, 0, 0)  
doc.save('ppa_report.docx') 
print("Text centered")

#
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph("The chart below clearly shows the students' considerable improvement in each skill from the Initial Assessment to the Post Program Assessment.")  
run = para.runs[0]
run.font.size = Pt(12)
doc.save('ppa_report.docx') 

#Graph insertion
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
import io
def rgb_to_normalized(rgb):
    return tuple([x / 255.0 for x in rgb])
background_color = rgb_to_normalized((32, 56, 100))  
bar1_color = rgb_to_normalized((197, 90, 17))  
bar2_color = rgb_to_normalized((146, 208, 80)) 
items = ['Linguistics', 'Problem Solving', 'Focus and Attention', 'Visual Processing', 'Memory']
post_scores = [linguisPPA, probsolPPA, focusatPPA, visualpPPA, memoryPPA] 
initial_scores = [linguisIA, probsolIA, focusatIA, visualpIA, memoryIA]  
y = np.arange(len(items))  
height = 0.3
fig, ax = plt.subplots(figsize=(9, 5)) 
plt.gcf().set_facecolor(background_color)
ax.set_facecolor(background_color)  
gap = 0.02
bar1 = ax.barh(y - height/2 - gap, initial_scores, height, label='Initial Assessment', color=bar1_color)  
bar2 = ax.barh(y + height/2 + gap, post_scores, height, label='Post Program Assessment', color=bar2_color) 
ax.set_title('Skill Growth - Average Skill Scores    Initial vs Post Program Assessment', color='white')
ax.set_yticks(y)
ax.set_yticklabels(items, color='white') 
ax.tick_params(axis='x', colors='white')  
ax.tick_params(axis='y', colors='white')  
ax.legend(facecolor='gray', edgecolor='white', loc='upper right', fontsize=10, framealpha=0.7)  
for i, rect in enumerate(bar1):
    width = rect.get_width()
    ax.text(width, rect.get_y() + rect.get_height() / 2, str(initial_scores[i]), ha='left', va='center', color='white')
for i, rect in enumerate(bar2):
    width = rect.get_width()
    ax.text(width, rect.get_y() + rect.get_height() / 2, str(post_scores[i]), ha='left', va='center', color='white')
ax.grid(True, which='both', axis='x' , linestyle='-', color='black', linewidth=0.6, alpha=0.2)
plt.tight_layout()
img_stream = io.BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0) 
doc = Document('ppa_report.docx')  
doc.add_paragraph("Here is the graph comparing Initial vs Post Program Assessment:")
table = doc.add_table(rows=1, cols=1)
cell = table.cell(0, 0)
cell.paragraphs[0].clear()  
cell.paragraphs[0].add_run().add_picture(img_stream, width=Inches(6)) 
cell.alignment = 1  
doc.save('ppa_report.docx')
img_stream.close() 
print("Graph inserted")

#next highlighted para
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph("Students have achieved remarkable strides, showcasing a remarkable 49% boost in Memory skills, closely followed by an impressive 32% improvement in Focus and Attention & Linguistics skills.")
run = para.runs[0]
run.font.size = Pt(12)
run.font.highlight_color = 7  
doc.save('ppa_report.docx') 

#
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('MEMORY')
para.alignment = WD_ALIGN_PARAGRAPH.CENTER  
run = para.runs[0]
run.bold = True  
run.font.size = Pt(18)  
run.font.color.rgb = RGBColor(0, 31, 95) 
run.font.underline = True  
doc.save('ppa_report.docx') 

# 
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('The table below captures the number of students in each score category/band at two specific points in the program – at the beginning, the')  
run = para.add_run(' IA')
run.bold = True  
para.add_run(', and at the end, the ')
run = para.add_run(' PPA.')
run.bold = True
for run in para.runs:
    run.font.size = Pt(12) 
doc.save('ppa_report.docx') 

#Table
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
doc = Document('ppa_report.docx')
table = doc.add_table(rows=6, cols=4)
table.cell(0, 0).text = 'Skill Score Range'
table.cell(0, 1).text = 'Description'
table.cell(0, 2).text = 'Number of students in IA'
table.cell(0, 3).text = 'Number of students in PPA'
table.cell(1, 0).text = '<=20 '
table.cell(1, 1).text = 'Rigorous and immediate intervention needed'
table.cell(1, 2).text = memIA20
table.cell(1, 3).text = memPPA20
table.cell(2, 0).text = '21-40'
table.cell(2, 1).text = 'Needs intervention at the earliest '
table.cell(2, 2).text = memIA40
table.cell(2, 3).text = memPPA40
table.cell(3, 0).text = '41-60'
table.cell(3, 1).text = 'Intervention will enhance academic performance'
table.cell(3, 2).text = memIA60
table.cell(3, 3).text = memPPA60
table.cell(4, 0).text = '61-80'
table.cell(4, 1).text = 'Can be groomed to be highly successful'
table.cell(4, 2).text = memIA80
table.cell(4, 3).text = memPPA80
table.cell(5, 0).text = '>80'
table.cell(5, 1).text = 'Prodigy – excellent performance! '
table.cell(5, 2).text = memIA100
table.cell(5, 3).text = memPPA100
for column in table.columns:
    for cell in column.cells:
        cell.width = Inches(1.5) 
for cell in table.rows[0].cells:
    for paragraph in cell.paragraphs:
        if paragraph.runs:  
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)
for row in table.rows:
    row.cells[0].width = Pt(70)  # First column narrow
    row.cells[1].width = Pt(250)  # Middle column wide
    row.cells[2].width = Pt(100)
    row.cells[3].width = Pt(100)
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  
            border.set(qn('w:sz'), '4')  
            border.set(qn('w:space'), '0') 
            border.set(qn('w:color'), '000000')  
            borders.append(border)
        tc_pr.append(borders)
row_colors = ['EE4646', 'EB795F', 'FAAC17', 'FDC350', '9ACC58']  
for i in range(6):  
    for j in range(2): 
        cell = table.cell(i, j)
        if i > 0: 
            cell_xml = cell._element
            tc_pr = cell_xml.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), row_colors[i-1] if i > 0 else 'FFFFFF') 
            tc_pr.append(shading)
para = doc.add_paragraph()
doc.save("ppa_report.docx") 
print("Table inserted")

#Graph insertion
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
import io
def rgb_to_normalized(rgb):
    return tuple([x / 255.0 for x in rgb])
background_color = rgb_to_normalized((255, 255, 255))  
bar1_color = rgb_to_normalized((255, 192, 0))  
bar2_color = rgb_to_normalized((146, 208, 80))  
items = ['<=20', '21-40', '41-60', '61-80', '>80']
initial_scores = [memgraIA1, memgraIA2, memgraIA3, memgraIA4, memgraIA5]
post_scores = [memgraPA1, memgraPA2, memgraPA3, memgraPA4, memgraPA5]
x = np.arange(len(items))
width = 0.25
fig, ax = plt.subplots(figsize=(9, 5))
plt.gcf().set_facecolor(background_color)  
ax.set_facecolor(background_color)  
gap = 0.02
bar1 = ax.bar(x - width/2 - gap, initial_scores, width, label='Initial Assessment', color=bar1_color)
bar2 = ax.bar(x + width/2 + gap, post_scores, width, label='Post Program Assessment', color=bar2_color)
ax.set_title('Memory - Skill Score Analysis', color='black',fontsize = 18, fontweight = 'bold')
ax.set_xticks(x)
ax.set_xticklabels(items, color='black', fontsize = 14)
ax.tick_params(axis='y', colors='black')
ax.tick_params(axis='x', colors='black')
ax.legend(facecolor='gray', edgecolor='white', loc='upper right', fontsize=12, framealpha=0.6)
for i, rect in enumerate(bar1):
    height = rect.get_height()
    ax.text(rect.get_x() + rect.get_width() / 2, height, str(initial_scores[i]), ha='center', va='bottom', color='black')
for i, rect in enumerate(bar2):
    height = rect.get_height()
    ax.text(rect.get_x() + rect.get_width() / 2, height, str(post_scores[i]), ha='center', va='bottom', color='black')
ax.grid(True, which='both', axis='y' , linestyle='-', color='black', linewidth=0.6, alpha=0.2)
plt.tight_layout()
img_stream = io.BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0)  
doc = Document('ppa_report.docx')
table = doc.add_table(rows=1, cols=1)
cell = table.cell(0, 0)
cell.paragraphs[0].clear()
cell.paragraphs[0].add_run().add_picture(img_stream, width=Inches(6))
cell.alignment = 1  
doc.save('ppa_report.docx')
img_stream.close() 
print("graph inserted")

# para
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph("At the commencement of the program, majority of students fell within the lower score bands. However, by the program's conclusion, a shift is seen, with a larger ratio of students moving into the higher score ranges. The 8% shift from initial to post assessment in >80 band indicates that in terms of MEMORY, the program has effectively helped a larger number of students in improving their skills.")  
run = para.runs[0]
run.font.size = Pt(12)
para = doc.add_paragraph()
doc.save('ppa_report.docx') 

# topic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('VISUAL PROCESSING ')
para.alignment = WD_ALIGN_PARAGRAPH.CENTER  
run = para.runs[0]
run.bold = True  
run.font.size = Pt(18)  
run.font.color.rgb = RGBColor(0, 31, 95) 
run.font.underline = True  
doc.save('ppa_report.docx') 

#
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('The table below captures the number of students in each score category/band at two specific points in the program – at the beginning, the')  
run = para.add_run(' IA')
run.bold = True  
para.add_run(', and at the end, the ')
run = para.add_run(' PPA.')
run.bold = True
for run in para.runs:
    run.font.size = Pt(12) 
doc.save('ppa_report.docx') 

# table
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
doc = Document('ppa_report.docx')
table = doc.add_table(rows=6, cols=4)
table.cell(0, 0).text = 'Skill Score Range'
table.cell(0, 1).text = 'Description'
table.cell(0, 2).text = 'Number of students in IA'
table.cell(0, 3).text = 'Number of students in PPA'
table.cell(1, 0).text = '<=20 '
table.cell(1, 1).text = 'Rigorous and immediate intervention needed'
table.cell(1, 2).text =  vispIA20
table.cell(1, 3).text =  vispPA20
table.cell(2, 0).text = '21-40'
table.cell(2, 1).text = 'Needs intervention at the earliest '
table.cell(2, 2).text =  vispIA40
table.cell(2, 3).text =  vispPA40
table.cell(3, 0).text = '41-60'
table.cell(3, 1).text = 'Intervention will enhance academic performance'
table.cell(3, 2).text =  vispIA60
table.cell(3, 3).text =  vispPA60
table.cell(4, 0).text = '61-80'
table.cell(4, 1).text = 'Can be groomed to be successfull'
table.cell(4, 2).text =  vispIA80
table.cell(4, 3).text =  vispPA80
table.cell(5, 0).text = '>80'
table.cell(5, 1).text = 'Prodigy – excellent performance!'
table.cell(5, 2).text =  vispIA81 
table.cell(5, 3).text =  vispPA81
for column in table.columns:
    for cell in column.cells:
        cell.width = Inches(1.5) 
for cell in table.rows[0].cells:
    for paragraph in cell.paragraphs:
        if paragraph.runs:  
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)
for row in table.rows:
    row.cells[0].width = Pt(70)  
    row.cells[1].width = Pt(250)  
    row.cells[2].width = Pt(100)
    row.cells[3].width = Pt(100)
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  
            border.set(qn('w:sz'), '4')  
            border.set(qn('w:space'), '0')  
            border.set(qn('w:color'), '000000')  
            borders.append(border)
        tc_pr.append(borders)
row_colors = ['EE4646', 'EB795F', 'FAAC17', 'FDC350', '9ACC58']  
for i in range(6):  # Iterate over all rows (index 0 to 5)
    for j in range(2):  # Apply colors only to the first two columns (index 0 and 1)
        cell = table.cell(i, j)
        if i > 0: 
            cell_xml = cell._element
            tc_pr = cell_xml.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), row_colors[i-1] if i > 0 else 'FFFFFF') 
            tc_pr.append(shading)
para = doc.add_paragraph()
doc.save("ppa_report.docx") 
print("Table inserted")

#graph
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
import io
def rgb_to_normalized(rgb):
    return tuple([x / 255.0 for x in rgb])
background_color = rgb_to_normalized((255, 255, 255))  
bar1_color = rgb_to_normalized((255, 192, 0))  
bar2_color = rgb_to_normalized((146, 208, 80))  
items = ['<=20', '21-40', '41-60', '61-80', '>80']
initial_scores = [visgraIA1, visgraIA2, visgraIA3, visgraIA4,visgraIA5]
post_scores = [visgraPA1, visgraPA2, visgraPA3, visgraPA4, visgraPA5]
x = np.arange(len(items))
width = 0.25
fig, ax = plt.subplots(figsize=(9, 5))
plt.gcf().set_facecolor(background_color)
ax.set_facecolor(background_color)  
gap = 0.02
bar1 = ax.bar(x - width/2 - gap, initial_scores, width, label='Initial Assessment', color=bar1_color)
bar2 = ax.bar(x + width/2 + gap, post_scores, width, label='Post Program Assessment', color=bar2_color)
ax.set_title('Visual Processing - Skill Score Analysis', color='black',fontsize = 18, fontweight = 'bold')
ax.set_xticks(x)
ax.set_xticklabels(items, color='black', fontsize = 14)
ax.tick_params(axis='y', colors='black')
ax.tick_params(axis='x', colors='black')
ax.legend(facecolor='gray', edgecolor='white', loc='upper right', fontsize=12, framealpha=0.6)
for i, rect in enumerate(bar1):
    height = rect.get_height()
    ax.text(rect.get_x() + rect.get_width() / 2, height, str(initial_scores[i]), ha='center', va='bottom', color='black')
for i, rect in enumerate(bar2):
    height = rect.get_height()
    ax.text(rect.get_x() + rect.get_width() / 2, height, str(post_scores[i]), ha='center', va='bottom', color='black')
ax.grid(True, which='both', axis='y' , linestyle='-', color='black', linewidth=0.6, alpha=0.2)
plt.tight_layout()
img_stream = io.BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0)  
doc = Document('ppa_report.docx')
table = doc.add_table(rows=1, cols=1)
cell = table.cell(0, 0)
cell.paragraphs[0].clear()
cell.paragraphs[0].add_run().add_picture(img_stream, width=Inches(6))
cell.alignment = 1  
doc.save('ppa_report.docx')
img_stream.close() 
print("graph inserted")

#text
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph("Where VISUAL PROCESSING is concerned, students have been able to move into the top bands with 17% improvement; this is a good development as it is bound to positively impact their thought processing as well.")  
run = para.runs[0]
run.font.size = Pt(12)
para = doc.add_paragraph()
doc.save('ppa_report.docx') 

#Topic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('FOCUS AND ATTENTION')
para.alignment = WD_ALIGN_PARAGRAPH.CENTER  
run = para.runs[0]
run.bold = True  
run.font.size = Pt(18)  
run.font.color.rgb = RGBColor(0, 31, 95) 
run.font.underline = True  
doc.save('ppa_report.docx') 

#
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('The table below captures the number of students in each score category/band at two specific points in the program – at the beginning, the')  
run = para.add_run(' IA')
run.bold = True  
para.add_run(', and at the end, the ')
run = para.add_run(' PPA.')
run.bold = True
for run in para.runs:
    run.font.size = Pt(12) 
doc.save('ppa_report.docx') 

# table
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
doc = Document('ppa_report.docx')
table = doc.add_table(rows=6, cols=4)
table.cell(0, 0).text = 'Skill Score Range'
table.cell(0, 1).text = 'Description'
table.cell(0, 2).text = 'Number of students in IA'
table.cell(0, 3).text = 'Number of students in PPA'
table.cell(1, 0).text = '<=20 '
table.cell(1, 1).text = 'Rigorous and immediate intervention needed'
table.cell(1, 2).text = focuatIA_20
table.cell(1, 3).text = focuatPA_20
table.cell(2, 0).text = '21-40'
table.cell(2, 1).text = 'Needs intervention at the earliest '
table.cell(2, 2).text = focuatIA_40
table.cell(2, 3).text = focuatPA_40 
table.cell(3, 0).text = '41-60'
table.cell(3, 1).text = 'Intervention will enhance academic performance'
table.cell(3, 2).text = focuatIA_60
table.cell(3, 3).text = focuatPA_60
table.cell(4, 0).text = '61-80'
table.cell(4, 1).text = 'Can be groomed to be highly successful'
table.cell(4, 2).text = focuatIA_80 
table.cell(4, 3).text = focuatPA_80
table.cell(5, 0).text = '>80'
table.cell(5, 1).text = 'Prodigy – excellent performance!'
table.cell(5, 2).text = focuatIA_100
table.cell(5, 3).text = focuatPA_100
for column in table.columns:
    for cell in column.cells:
        cell.width = Inches(1.5) 
for cell in table.rows[0].cells:
    for paragraph in cell.paragraphs:
        if paragraph.runs:  
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)
for row in table.rows:
    row.cells[0].width = Pt(70)   
    row.cells[1].width = Pt(250)   
    row.cells[2].width = Pt(100)
    row.cells[3].width = Pt(100)
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')   
            border.set(qn('w:sz'), '4')   
            border.set(qn('w:space'), '0')   
            border.set(qn('w:color'), '000000')   
            borders.append(border)
        tc_pr.append(borders)
row_colors = ['EE4646', 'EB795F', 'FAAC17', 'FDC350', '9ACC58']  
for i in range(6):   
    for j in range(2):  
        cell = table.cell(i, j)
        if i > 0: 
            cell_xml = cell._element
            tc_pr = cell_xml.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), row_colors[i-1] if i > 0 else 'FFFFFF') 
            tc_pr.append(shading)
para = doc.add_paragraph()
doc.save("ppa_report.docx") 
print("Table inserted")

#graph
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
import io
def rgb_to_normalized(rgb):
    return tuple([x / 255.0 for x in rgb])
background_color = rgb_to_normalized((255, 255, 255))  
bar1_color = rgb_to_normalized((255, 192, 0))  
bar2_color = rgb_to_normalized((146, 208, 80))  
items = ['<=20', '21-40', '41-60', '61-80', '>80']
initial_scores = [focgraIA1,focgraIA2, focgraIA3, focgraIA4, focgraIA5]
post_scores = [focgraPA1, focgraPA2, focgraPA3, focgraPA4, focgraPA5]
x = np.arange(len(items))
width = 0.25
fig, ax = plt.subplots(figsize=(9, 5))
plt.gcf().set_facecolor(background_color)
ax.set_facecolor(background_color)  
gap = 0.02
bar1 = ax.bar(x - width/2 - gap, initial_scores, width, label='Initial Assessment', color=bar1_color)
bar2 = ax.bar(x + width/2 + gap, post_scores, width, label='Post Program Assessment', color=bar2_color)
ax.set_title('Focus and Attention - Skill Score Analysis', color='black',fontsize = 18, fontweight = 'bold')
ax.set_xticks(x)
ax.set_xticklabels(items, color='black', fontsize = 14)
ax.tick_params(axis='y', colors='black')
ax.tick_params(axis='x', colors='black')
ax.legend(facecolor='gray', edgecolor='white', loc='upper right', fontsize=12, framealpha=0.6)
for i, rect in enumerate(bar1):
    height = rect.get_height()
    ax.text(rect.get_x() + rect.get_width() / 2, height, str(initial_scores[i]), ha='center', va='bottom', color='black')
for i, rect in enumerate(bar2):
    height = rect.get_height()
    ax.text(rect.get_x() + rect.get_width() / 2, height, str(post_scores[i]), ha='center', va='bottom', color='black')
ax.grid(True, which='both', axis='y' , linestyle='-', color='black', linewidth=0.6, alpha=0.2)
plt.tight_layout()
img_stream = io.BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0)  
doc = Document('ppa_report.docx')
table = doc.add_table(rows=1, cols=1)
cell = table.cell(0, 0)
cell.paragraphs[0].clear()
cell.paragraphs[0].add_run().add_picture(img_stream, width=Inches(6))
cell.alignment = 1  
doc.save('ppa_report.docx')
img_stream.close() 
print("Graph inserted")

#text
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph("In terms of FOCUS AND ATTENTION, students have successfully advanced into the top band with 26% improvement marking a notable enhancement that is likely to enhance their cognitive processing positively.")  
run = para.runs[0]
run.font.size = Pt(12)
para = doc.add_paragraph()
doc.save('ppa_report.docx') 

#Topic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('Problem Solving')
para.alignment = WD_ALIGN_PARAGRAPH.CENTER  
run = para.runs[0]
run.bold = True  
run.font.size = Pt(18)  
run.font.color.rgb = RGBColor(0, 31, 95) 
run.font.underline = True  
doc.save('ppa_report.docx') 

#
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('The table below captures the number of students in each score category/band at two specific points in the program – at the beginning, the')  
run = para.add_run(' IA')
run.bold = True  
para.add_run(', and at the end, the ')
run = para.add_run(' PPA.')
run.bold = True
for run in para.runs:
    run.font.size = Pt(12) 
doc.save('ppa_report.docx') 

# table
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
doc = Document('ppa_report.docx')
table = doc.add_table(rows=6, cols=4)
table.cell(0, 0).text = 'Skill Score Range'
table.cell(0, 1).text = 'Description'
table.cell(0, 2).text = 'Number of students in IA'
table.cell(0, 3).text = 'Number of students in PPA'
table.cell(1, 0).text = '<=20 '
table.cell(1, 1).text = 'Rigorous and immediate intervention needed'
table.cell(1, 2).text = prosolIA_20 
table.cell(1, 3).text = prosolPA_20
table.cell(2, 0).text = '21-40'
table.cell(2, 1).text = 'Needs intervention at the earliest '
table.cell(2, 2).text = prosolIA_40
table.cell(2, 3).text = prosolPA_40
table.cell(3, 0).text = '41-60'
table.cell(3, 1).text = 'Intervention will enhance academic performance'
table.cell(3, 2).text = prosolIA_60
table.cell(3, 3).text = prosolPA_60
table.cell(4, 0).text = '61-80'
table.cell(4, 1).text = 'Can be groomed to be highly successful'
table.cell(4, 2).text = prosolIA_80
table.cell(4, 3).text = prosolPA_80
table.cell(5, 0).text = '>80'
table.cell(5, 1).text = 'Prodigy – excellent performance!'
table.cell(5, 2).text = prosolIA_100
table.cell(5, 3).text = prosolPA_100
for column in table.columns:
    for cell in column.cells:
        cell.width = Inches(1.5) 
for cell in table.rows[0].cells:
    for paragraph in cell.paragraphs:
        if paragraph.runs:  
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)
for row in table.rows:
    row.cells[0].width = Pt(70)  # First column narrow
    row.cells[1].width = Pt(250)  # Middle column wide
    row.cells[2].width = Pt(100)
    row.cells[3].width = Pt(100)
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  # Solid border
            border.set(qn('w:sz'), '4')  # Border width
            border.set(qn('w:space'), '0')  # No space between cells
            border.set(qn('w:color'), '000000')  # Black color for borders
            borders.append(border)
        tc_pr.append(borders)
row_colors = ['EE4646', 'EB795F', 'FAAC17', 'FDC350', '9ACC58']  
for i in range(6):  # Iterate over all rows (index 0 to 5)
    for j in range(2):  # Apply colors only to the first two columns (index 0 and 1)
        cell = table.cell(i, j)
        if i > 0: 
            cell_xml = cell._element
            tc_pr = cell_xml.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), row_colors[i-1] if i > 0 else 'FFFFFF') 
            tc_pr.append(shading)
para = doc.add_paragraph()
doc.save("ppa_report.docx") 
print("Table inserted")

#graph
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
import io
def rgb_to_normalized(rgb):
    return tuple([x / 255.0 for x in rgb])
background_color = rgb_to_normalized((255, 255, 255))  # White background
bar1_color = rgb_to_normalized((255, 192, 0))  
bar2_color = rgb_to_normalized((146, 208, 80))  
items = ['<=20', '21-40', '41-60', '61-80', '>80']
initial_scores = [prograIA1, prograIA2, prograIA3, prograIA4, prograIA5]
post_scores = [prograPA1, prograPA2, prograPA3, prograPA4, prograPA5]
x = np.arange(len(items))
width = 0.25
fig, ax = plt.subplots(figsize=(9, 5))
plt.gcf().set_facecolor(background_color)
ax.set_facecolor(background_color)  
gap = 0.02
bar1 = ax.bar(x - width/2 - gap, initial_scores, width, label='Initial Assessment', color=bar1_color)
bar2 = ax.bar(x + width/2 + gap, post_scores, width, label='Post Program Assessment', color=bar2_color)
ax.set_title('Problem Solving- Skill Score Analysis', color='black',fontsize = 18, fontweight = 'bold')
ax.set_xticks(x)
ax.set_xticklabels(items, color='black', fontsize = 14)
ax.tick_params(axis='y', colors='black')
ax.tick_params(axis='x', colors='black')
ax.legend(facecolor='gray', edgecolor='white', loc='upper right', fontsize=12, framealpha=0.6)
for i, rect in enumerate(bar1):
    height = rect.get_height()
    ax.text(rect.get_x() + rect.get_width() / 2, height, str(initial_scores[i]), ha='center', va='bottom', color='black')
for i, rect in enumerate(bar2):
    height = rect.get_height()
    ax.text(rect.get_x() + rect.get_width() / 2, height, str(post_scores[i]), ha='center', va='bottom', color='black')
ax.grid(True, which='both', axis='y' , linestyle='-', color='black', linewidth=0.6, alpha=0.2)
plt.tight_layout()
img_stream = io.BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0)  
doc = Document('ppa_report.docx')
table = doc.add_table(rows=1, cols=1)
cell = table.cell(0, 0)
cell.paragraphs[0].clear()
cell.paragraphs[0].add_run().add_picture(img_stream, width=Inches(6))
cell.alignment = 1  
doc.save('ppa_report.docx')
img_stream.close() 
print("graph inserted")

#text
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph("Students have achieved top bands in PROBLEM SOLVING, signifying substantial progress likely to enhance their cognitive processing positively. Since this involves analytical skills and is heavy on thinking out of the box, it is necessary to work on it harder.")  
run = para.runs[0]
run.font.size = Pt(12)
para = doc.add_paragraph()
doc.save('ppa_report.docx') 

#Topic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('LINGUISTICS')
para.alignment = WD_ALIGN_PARAGRAPH.CENTER  
run = para.runs[0]
run.bold = True  
run.font.size = Pt(18)  
run.font.color.rgb = RGBColor(0, 31, 95) 
run.font.underline = True  
doc.save('ppa_report.docx') 

#
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('The table below captures the number of students in each score category/band at two specific points in the program – at the beginning, the')  
run = para.add_run(' IA')
run.bold = True  
para.add_run(', and at the end, the ')
run = para.add_run(' PPA.')
run.bold = True
for run in para.runs:
    run.font.size = Pt(12) 
doc.save('ppa_report.docx') 

# table
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
doc = Document('ppa_report.docx')
table = doc.add_table(rows=6, cols=4)
table.cell(0, 0).text = 'Skill Score Range'
table.cell(0, 1).text = 'Description'
table.cell(0, 2).text = 'Number of students in IA'
table.cell(0, 3).text = 'Number of students in PPA'
table.cell(1, 0).text = '<=20 '
table.cell(1, 1).text = 'Rigorous and immediate intervention needed'
table.cell(1, 2).text = linguiIA_20
table.cell(1, 3).text = linguiPA_20
table.cell(2, 0).text = '21-40'
table.cell(2, 1).text = 'Needs intervention at the earliest '
table.cell(2, 2).text = linguiIA_40
table.cell(2, 3).text = linguiPA_40
table.cell(3, 0).text = '41-60'
table.cell(3, 1).text = 'Intervention will enhance academic performance'
table.cell(3, 2).text = linguiIA_60
table.cell(3, 3).text = linguiPA_60
table.cell(4, 0).text = '61-80'
table.cell(4, 1).text = 'Can be groomed to be highly successful'
table.cell(4, 2).text = linguiIA_80
table.cell(4, 3).text = linguiPA_80
table.cell(5, 0).text = '>80'
table.cell(5, 1).text = 'Prodigy – excellent performance!'
table.cell(5, 2).text = linguiIA_100
table.cell(5, 3).text = linguiPA_100
for column in table.columns:
    for cell in column.cells:
        cell.width = Inches(1.5) 
for cell in table.rows[0].cells:
    for paragraph in cell.paragraphs:
        if paragraph.runs:  
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)
for row in table.rows:
    row.cells[0].width = Pt(70)  
    row.cells[1].width = Pt(250)  
    row.cells[2].width = Pt(100)
    row.cells[3].width = Pt(100)
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  
            border.set(qn('w:sz'), '4')  
            border.set(qn('w:space'), '0')  
            border.set(qn('w:color'), '000000')  
            borders.append(border)
        tc_pr.append(borders)
row_colors = ['EE4646', 'EB795F', 'FAAC17', 'FDC350', '9ACC58']  
for i in range(6):  
    for j in range(2):  
        cell = table.cell(i, j)
        if i > 0: 
            cell_xml = cell._element
            tc_pr = cell_xml.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), row_colors[i-1] if i > 0 else 'FFFFFF') 
            tc_pr.append(shading)
para = doc.add_paragraph()
doc.save("ppa_report.docx") 
print("Table inserted")

#graph
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
import io 
def rgb_to_normalized(rgb):
    return tuple([x / 255.0 for x in rgb])
background_color = rgb_to_normalized((255, 255, 255))  
bar1_color = rgb_to_normalized((255, 192, 0))  
bar2_color = rgb_to_normalized((146, 208, 80))  
items = ['<=20', '21-40', '41-60', '61-80', '>80']
initial_scores = [lingraIA1, lingraIA1, lingraIA3, lingraIA4, lingraIA5]
post_scores = [lingraPA1, lingraPA2, lingraPA3, lingraPA4, lingraPA5]
x = np.arange(len(items))
width = 0.25
fig, ax = plt.subplots(figsize=(9, 5))
plt.gcf().set_facecolor(background_color)
ax.set_facecolor(background_color)  
gap = 0.02
bar1 = ax.bar(x - width/2 - gap, initial_scores, width, label='Initial Assessment', color=bar1_color)
bar2 = ax.bar(x + width/2 + gap, post_scores, width, label='Post Program Assessment', color=bar2_color)
ax.set_title('Linguistics - Skill Score Analysis', color='black',fontsize = 18, fontweight = 'bold')
ax.set_xticks(x)
ax.set_xticklabels(items, color='black', fontsize = 14)
ax.tick_params(axis='y', colors='black')
ax.tick_params(axis='x', colors='black')
ax.legend(facecolor='gray', edgecolor='white', loc='upper right', fontsize=12, framealpha=0.6)
for i, rect in enumerate(bar1):
    height = rect.get_height()
    ax.text(rect.get_x() + rect.get_width() / 2, height, str(initial_scores[i]), ha='center', va='bottom', color='black')
for i, rect in enumerate(bar2):
    height = rect.get_height()
    ax.text(rect.get_x() + rect.get_width() / 2, height, str(post_scores[i]), ha='center', va='bottom', color='black')
ax.grid(True, which='both', axis='y' , linestyle='-', color='black', linewidth=0.6, alpha=0.2)
plt.tight_layout()
img_stream = io.BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0)  
doc = Document('ppa_report.docx')
table = doc.add_table(rows=1, cols=1)
cell = table.cell(0, 0)
cell.paragraphs[0].clear()
cell.paragraphs[0].add_run().add_picture(img_stream, width=Inches(6))
cell.alignment = 1  
doc.save('ppa_report.docx')
img_stream.close() 
print("Graph inserted")

#text
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph("Recognizing the crucial role of LINGUISTICS in bolstering other learning domains is imperative for its enhancement. While students are performing satisfactorily in this skill, there remains room for improvement.")  
run = para.runs[0]
run.font.size = Pt(12)
para = doc.add_paragraph()
doc.save('ppa_report.docx')

# text topic
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph('5 Do You Know How Well They Did? ')
run = para.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(45, 116, 181)
doc.save('ppa_report.docx')

#
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph("The table below indicates the scores of students from two categories – those who have completed less than eight sessions in CAUP and those who have completed more than eight sessions in CAUP. ")  
para = doc.add_paragraph("The total number of students who have taken ")
run = para.add_run(' IA')
run.bold = True  
para.add_run(' &')
run = para.add_run(' PPA')
run.bold = True
para.add_run(' is ')
run = para.add_run( noofstudentsdoneppa)
run.bold = True
run = para.runs[0]
run.font.size = Pt(12)
para = doc.add_paragraph()
doc.save('ppa_report.docx')
#Table Creation
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
doc = Document('ppa_report.docx')
table = doc.add_table(rows=7, cols=7)
headers = ['S.No', 'Grade', 'Total students completed PPA ', 'Number of students who have taken below 8 sessions in CAUP ', 'Average score of PPA taken by students who have completed below 8 sessions in CAUP ', 'Number of students who have taken 8 and more sessions in CAUP ', 'Average score of PPA taken by students who have taken 8 and above session we in CAUP']
rows_data = [
    ['1', 'Grade III', '63', '0', grade3avgscbel8, '63', grade3avgsabove8],
    ['2', 'Grade IV', '58', '0', grade4avgscbel8, '58', grade4avgsabove8],
    ['3', 'Grade V', '69', '0', grade5avgscbel8, '69', grade5avgsabove8],
    ['4', 'Grade VI', '58', '0', grade6avgscbel8, '58', grade6avgsabove8],
    ['5', 'Grade VII', '36', '0', grade7avgscbel8, '36', grade7avgsabove8],
    ['6', 'Grade VIII', '45', '0', grade8avgscbel8, '45', grade8avgsabove8]
]
for col, header in enumerate(headers):
    cell = table.cell(0, col)
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(14)  
    cell_xml = cell._element
    tc_pr = cell_xml.get_or_add_tcPr()
    shading = OxmlElement('w:shd')  # Create shading element for background color
    shading.set(qn('w:fill'), '92CDEC')  # Set background color (example: light green RGB)
    tc_pr.append(shading)
    cell_text = cell._element
    tc_pr = cell_text.get_or_add_tcPr()
    rotation = OxmlElement('w:textDirection')
    rotation.set(qn('w:val'), 'btLr') 
    tc_pr.append(rotation)
    if col in [0, 2, 3, 5]:  
        cell_width = Inches(1.5)  # Adjust to a smaller width
    else: 
        cell_width = Inches(2.5)  
    tc_pr = cell_xml.get_or_add_tcPr()
    tc_w = OxmlElement('w:tcW')
    tc_w.set(qn('w:w'), str(int(cell_width * 1440)))  # Convert inches to twips (1 inch = 1440 twips)
    tc_w.set(qn('w:type'), 'dxa')  # Type 'dxa' means it's in twips
    tc_pr.append(tc_w)
table.rows[0].height = Inches(2.5) 
for col in range(len(headers)):
    cell = table.cell(0, col)
    cell_xml = cell._element
    tc_pr = cell_xml.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(4))  # Adjust the factor as needed
        border.set(qn('w:space'), '0')  # No space between borders
        border.set(qn('w:color'), '000000')  # Black border color
        borders.append(border)
    tc_pr.append(borders)
for i, row_data in enumerate(rows_data, start=1):
    for j, text in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = text
        if j in [4, 6]:  
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(12) 
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4') 
            border.set(qn('w:space'), '0')  
            border.set(qn('w:color'), '000000')  
            borders.append(border)
        tc_pr.append(borders)
doc.save('ppa_report.docx')
print("Table inserted")

#
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph('It is evident that students who are regular with CAUP sessions and complete all the required/scheduled sessions have a brighter chance of improving their skills.')
run = para.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 175, 80)
doc.save('ppa_report.docx')

#Logo
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document('ppa_report.docx')  
para = doc.add_paragraph()
run = para.add_run()
logo_path = r'C:\Users\varsh\OneDrive\Sathyabama\Edsix\New folder\logo_image.png'  
run.add_picture(logo_path, width=Inches(2)) 
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save('ppa_report.docx')  

# topic
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('6 Sub-Skill Score Analysis')
run = para.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(45, 116, 181)
doc.save('ppa_report.docx')

# text
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph("The table below shows the average sub skill scores. The scores were taken based on the CAUP sessions played throughout the program.")  
run = para.runs[0]
run.font.size = Pt(12)
para = doc.add_paragraph()
doc.save('ppa_report.docx')

# Graph3 
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
import io
def rgb_to_normalized(rgb):
    return tuple([x / 255.0 for x in rgb])
rgb = (146, 208, 80)  
normalized_rgb = rgb_to_normalized(rgb)
categories = ['Contextual Recall', 'Associative Recall', 'Working Memory',  'Spatial Awareness', 'Conservation', 'Creative Thinking-Visualisation',
              'Creative Thinking-Synthesis', 'Sustaining attention', 'Selective attention', 'Divided attention', 'Inductive Reasoning', 'Deductive Reasoning',
              'Classification', 'Distinguishing', 'Pattern Recognition/Sequencing', 'Inferring', 'Prediction and Conclusion', 'Critical Thinking-Visualisation', 
              'Language Processing', 'Assimilation', 'Accomodation']
scores = [subskillgraph1, subskillgraph2, subskillgraph3, subskillgraph4, subskillgraph5, subskillgraph6, subskillgraph7, subskillgraph8, subskillgraph9, subskillgraph10, subskillgraph11, subskillgraph12, subskillgraph13, subskillgraph14, subskillgraph15, subskillgraph16, subskillgraph17, subskillgraph18, subskillgraph19, subskillgraph20, subskillgraph21]
graph_background_rgb = (32, 56, 100)  
axes_background_rgb = (32, 56, 100)  
graph_background = rgb_to_normalized(graph_background_rgb)
axes_background = rgb_to_normalized(axes_background_rgb)
plt.figure(figsize=(9, 6))
plt.gcf().set_facecolor(graph_background)
ax = plt.gca()
ax.set_facecolor(axes_background)
bar_width = 0.4 
bars = plt.bar(categories, scores, color=normalized_rgb, width=bar_width)
plt.title('Sub Skill Average Scores', fontsize=16, fontweight='bold', color='white')
plt.xticks(rotation=45, ha='right', fontsize=10, color='white')
plt.yticks(np.arange(0, 90, 10), color='white')
ax.grid(True, which='both', axis='y', linestyle='-', linewidth=0.5, color='black')
for i, bar in enumerate(bars):
    height = bar.get_height()  
    plt.text(bar.get_x() + bar.get_width() / 2, height + 1, f'{height:.2f}', ha='center', va='bottom', fontsize=10, color='white')
plt.tight_layout()
img_stream = io.BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0)  
doc = Document('ppa_report.docx')
table = doc.add_table(rows=1, cols=1)
cell = table.cell(0, 0)
cell.paragraphs[0].clear() 
cell.paragraphs[0].add_run().add_picture(img_stream, width=Inches(6)) 
doc.save('ppa_report.docx')
img_stream.close()
print("Graph inserted")

# Text
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph("These scores offer a comprehensive gauge of the participants' proficiency across various subskills, offering valuable insights into both strengths and areas for improvement.")  
para = doc.add_paragraph("By dissecting these scores, we can uncover patterns and trends that enrich our understanding of the participants' skill development journey throughout the CAUP sessions.")
run = para.runs[0]
run.font.size = Pt(12)
para = doc.add_paragraph()
doc.save('ppa_report.docx')

# topic
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('7 Achievers in PPA 2023-23')
run = para.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(45, 116, 181)
doc.save('ppa_report.docx')

# logo
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document('ppa_report.docx')  
para = doc.add_paragraph()
run = para.add_run()
logo_path = r'C:\Users\varsh\OneDrive\Sathyabama\Edsix\New folder\logo5.png'  
run.add_picture(logo_path, width=Inches(7)) 
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save('ppa_report.docx')  
print("Logo added")

# topic
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('7.1 BSPI Toppers by Grade & Section:')
run = para.runs[0]
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(46, 116, 181)
doc.save('ppa_report.docx')

#
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph("The table below recognizes the students who have exhibited the best overall performance of Grades ")  
run = para.add_run(' 3,4,5,6,7,8')
run.bold = True  
para.add_run(' during the')
run = para.add_run(' PPA.')
run.bold = True
para = doc.add_paragraph("This has been computed by comparing the BSPI of students within each section of each grade.")
run = para.runs[0]
run.font.size = Pt(12)
doc.save('ppa_report.docx')

# Table
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
doc = Document('ppa_report.docx')
table = doc.add_table(rows=12, cols=4)
background_rgb = RGBColor(255, 192, 0)
text_content = [
    ['Student Name ', 'Grade', 'Section', 'Score'],
    [bspitop_studname1, bspitop_grade1, bspitop_sec1, bspitop_bspi1],
    [bspitop_studname2, bspitop_grade2, bspitop_sec2, bspitop_bspi2],
    [bspitop_studname3, bspitop_grade3, bspitop_sec3, bspitop_bspi3],
    [' ', ' ', ' ', ' '],
    ['KESAV', 'Row 6 Col 2', 'Row 6 Col 3', 'Row 6 Col 4'],
    ['THISANTH', 'Row 7 Col 2', 'Row 7 Col 3', 'Row 7 Col 4'],
    ['AKSHAYA KEERTHI ', 'Row 8 Col 2', 'Row 8 Col 3', 'Row 8 Col 4'],
    ['VIDHARSANA ', 'Row 9 Col 2', 'Row 9 Col 3', 'Row 9 Col 4'],
    ['SAHAANA', 'Row 10 Col 2', 'Row 10 Col 3', 'Row 10 Col 4'],
    ['ROSHNA ', 'Row 11 Col 2', 'Row 11 Col 3', 'Row 11 Col 4'],
    ['KALEESHWARAN', 'Row 12 Col 2', 'Row 12 Col 3', 'Row 12 Col 4']
]
for row_idx in range(12):
    for col_idx in range(4):
        cell = table.cell(row_idx, col_idx)
        cell.text = text_content[row_idx][col_idx]
        if row_idx == 0:
            cell_xml = cell._element
            tc_pr = cell_xml.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), f'{background_rgb[0]:02X}{background_rgb[1]:02X}{background_rgb[2]:02X}')
            tc_pr.append(shading)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
for row in table.rows:
    row.cells[0].width = Pt(200)  # First column narrow
    row.cells[1].width = Pt(80)  # Middle column wide
    row.cells[2].width = Pt(80)
    row.cells[3].width = Pt(80)
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  # Set border style
            border.set(qn('w:sz'), '4')  # Set border thickness
            border.set(qn('w:space'), '0')  # No space between borders
            border.set(qn('w:color'), '000000')  # Set border color to black
            borders.append(border)
        tc_pr.append(borders)
doc.save("ppa_report.docx")
print("Table inserted")

#
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('7.2 Skill Toppers by Grade & Section: ')
run = para.runs[0]  
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(46, 116, 181)
doc.save('ppa_report.docx')

#
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph("In addition to the BSPI, a cumulative average of scores across five cognitive skills. They have shown proficiency in individual skills of")  
run = para.add_run(' Grades 3,4,5,6,7,8')
run.bold = True
para.add_run(' during the ')
run = para.add_run(' PPA.')
run.bold = True
run = para.runs[0]
run.font.size = Pt(12)
doc.save('ppa_report.docx')

#
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph('Memory Toppers')
run = para.runs[0]
run.bold = True
run.font.size = Pt(12)
doc.save('ppa_report.docx')

# Table
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
doc = Document('ppa_report.docx')
table = doc.add_table(rows=12, cols=4)
background_rgb = RGBColor(255, 192, 0)
text_content = [
    ['Student Name ', 'Grade', 'Section', 'Score'],
    [memtoppers_name1, memtoppers_grade1, memtoppers_sec1, memtoppers_score1],
    [memtoppers_name2, memtoppers_grade2, memtoppers_sec2, memtoppers_score2],
    [memtoppers_name3, memtoppers_grade3, memtoppers_sec3, memtoppers_score3],
    [' ', ' ', ' ', ' '],
    [' ', ' ', ' ', ' '],
    [' ', 'Row 7 Col 2', 'Row 7 Col 3', 'Row 7 Col 4'],
    [' ', 'Row 8 Col 2', 'Row 8 Col 3', 'Row 8 Col 4'],
    [' ', 'Row 9 Col 2', 'Row 9 Col 3', 'Row 9 Col 4'],
    [' ', 'Row 10 Col 2', 'Row 10 Col 3', 'Row 10 Col 4'],
    [' ', 'Row 11 Col 2', 'Row 11 Col 3', 'Row 11 Col 4'],
    [' ', 'Row 12 Col 2', 'Row 12 Col 3', 'Row 12 Col 4'],
    [' ', 'Row 11 Col 2', 'Row 11 Col 3', 'Row 11 Col 4'],
    [' ', 'Row 11 Col 2', 'Row 11 Col 3', 'Row 11 Col 4'],
    [' ', 'Row 11 Col 2', 'Row 11 Col 3', 'Row 11 Col 4'],
]
for row_idx in range(12):
    for col_idx in range(4):
        cell = table.cell(row_idx, col_idx)
        cell.text = text_content[row_idx][col_idx]
        if row_idx == 0:
            cell_xml = cell._element
            tc_pr = cell_xml.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), f'{background_rgb[0]:02X}{background_rgb[1]:02X}{background_rgb[2]:02X}')
            tc_pr.append(shading)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
for row in table.rows:
    row.cells[0].width = Pt(200)  # First column narrow
    row.cells[1].width = Pt(80)  # Middle column wide
    row.cells[2].width = Pt(80)
    row.cells[3].width = Pt(80)
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  # Set border style
            border.set(qn('w:sz'), '4')  # Set border thickness
            border.set(qn('w:space'), '0')  # No space between borders
            border.set(qn('w:color'), '000000')  # Set border color to black
            borders.append(border)
        tc_pr.append(borders)
doc.save("ppa_report.docx")
print("Table inserted")

#
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph()
para = doc.add_paragraph('Visual Processing Toppers:')
run = para.runs[0]
run.bold = True
run.font.size = Pt(12)
doc.save('ppa_report.docx')

# table
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
doc = Document('ppa_report.docx')
table = doc.add_table(rows=12, cols=4)
background_rgb = RGBColor(255, 192, 0)
text_content = [
    ['Student Name ', 'Grade', 'Section', 'Score'],
    [vptoppers_name1, vptoppers_grade1, vptoppers_sec1, vptoppers_score1],
    [vptoppers_name2, vptoppers_grade2, vptoppers_sec2, vptoppers_score2],
    [vptoppers_name3, vptoppers_grade3, vptoppers_sec3, vptoppers_score3],
    [' ', ' ', ' ', ' '],
    [' ', ' ', ' ', ' '],
    [' ', 'Row 7 Col 2', 'Row 7 Col 3', 'Row 7 Col 4'],
    [' ', 'Row 8 Col 2', 'Row 8 Col 3', 'Row 8 Col 4'],
    [' ', 'Row 9 Col 2', 'Row 9 Col 3', 'Row 9 Col 4'],
    [' ', 'Row 10 Col 2', 'Row 10 Col 3', 'Row 10 Col 4'],
    [' ', 'Row 11 Col 2', 'Row 11 Col 3', 'Row 11 Col 4'],
    [' ', 'Row 12 Col 2', 'Row 12 Col 3', 'Row 12 Col 4'],
    [' ', 'Row 11 Col 2', 'Row 11 Col 3', 'Row 11 Col 4'],
    [' ', 'Row 11 Col 2', 'Row 11 Col 3', 'Row 11 Col 4'],
    [' ', 'Row 11 Col 2', 'Row 11 Col 3', 'Row 11 Col 4'],
]
for row_idx in range(12):
    for col_idx in range(4):
        cell = table.cell(row_idx, col_idx)
        cell.text = text_content[row_idx][col_idx]
        if row_idx == 0:
            cell_xml = cell._element
            tc_pr = cell_xml.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), f'{background_rgb[0]:02X}{background_rgb[1]:02X}{background_rgb[2]:02X}')
            tc_pr.append(shading)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
for row in table.rows:
    row.cells[0].width = Pt(200)  # First column narrow
    row.cells[1].width = Pt(80)  # Middle column wide
    row.cells[2].width = Pt(80)
    row.cells[3].width = Pt(80)
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  # Set border style
            border.set(qn('w:sz'), '4')  # Set border thickness
            border.set(qn('w:space'), '0')  # No space between borders
            border.set(qn('w:color'), '000000')  # Set border color to black
            borders.append(border)
        tc_pr.append(borders)
doc.save("ppa_report.docx")
print("Table inserted")

#
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph('Focus and Attention Toppers:')
run = para.runs[0]
run.bold = True
run.font.size = Pt(12)
doc.save('ppa_report.docx')

#Table
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
doc = Document('ppa_report.docx')
table = doc.add_table(rows=12, cols=4)
background_rgb = RGBColor(255, 192, 0)
text_content = [
    ['Student Name ', 'Grade', 'Section', 'Score'],
    [FAtoppers_name1, FAtoppers_grade1, FAtoppers_sec1, FAtoppers_score1],
    [FAtoppers_name2, FAtoppers_grade2, FAtoppers_sec2, FAtoppers_score2],
    [FAtoppers_name3, FAtoppers_grade3, FAtoppers_sec3, FAtoppers_score3],
    [' ', ' ', ' ', ' '],
    [' ', ' ', ' ', ' '],
    [' ', 'Row 7 Col 2', 'Row 7 Col 3', 'Row 7 Col 4'],
    [' ', 'Row 8 Col 2', 'Row 8 Col 3', 'Row 8 Col 4'],
    [' ', 'Row 9 Col 2', 'Row 9 Col 3', 'Row 9 Col 4'],
    [' ', 'Row 10 Col 2', 'Row 10 Col 3', 'Row 10 Col 4'],
    [' ', 'Row 11 Col 2', 'Row 11 Col 3', 'Row 11 Col 4'],
    [' ', 'Row 12 Col 2', 'Row 12 Col 3', 'Row 12 Col 4'],
    [' ', 'Row 11 Col 2', 'Row 11 Col 3', 'Row 11 Col 4'],
    [' ', 'Row 11 Col 2', 'Row 11 Col 3', 'Row 11 Col 4'],
    [' ', 'Row 11 Col 2', 'Row 11 Col 3', 'Row 11 Col 4'],
]
for row_idx in range(12):
    for col_idx in range(4):
        cell = table.cell(row_idx, col_idx)
        cell.text = text_content[row_idx][col_idx]
        if row_idx == 0:
            cell_xml = cell._element
            tc_pr = cell_xml.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), f'{background_rgb[0]:02X}{background_rgb[1]:02X}{background_rgb[2]:02X}')
            tc_pr.append(shading)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
for row in table.rows:
    row.cells[0].width = Pt(200) 
    row.cells[1].width = Pt(80)  
    row.cells[2].width = Pt(80)
    row.cells[3].width = Pt(80)
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  # Set border style
            border.set(qn('w:sz'), '4')  # Set border thickness
            border.set(qn('w:space'), '0')  # No space between borders
            border.set(qn('w:color'), '000000')  # Set border color to black
            borders.append(border)
        tc_pr.append(borders)
doc.save("ppa_report.docx")
print("Table inserted")

# 
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph('Problem Solving Toppers')
run = para.runs[0]
run.bold = True
run.font.size = Pt(12)
doc.save('ppa_report.docx')

# Table
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
doc = Document('ppa_report.docx')
table = doc.add_table(rows=12, cols=4)
background_rgb = RGBColor(255, 192, 0)
text_content = [
    ['Student Name ', 'Grade', 'Section', 'Score'],
    [PSoltoppers_name1, PSoltoppers_grade1, PSoltoppers_sec1, PSoltoppers_score1],
    [PSoltoppers_name2, PSoltoppers_grade2, PSoltoppers_sec2, PSoltoppers_score2],
    [PSoltoppers_name3, PSoltoppers_grade3, PSoltoppers_sec3, PSoltoppers_score3],
    ['KIRAN', '4', 'B', '91'],
    ['KESAV', '5', 'A', '99'],
    ['KRISHNA KANNIKA ', '5', 'B', '90'],
    ['KALKI MUTHAMIL', '6', 'A', '85'],
    ['NIVEDHA', '6', 'B', '89'],
    ['AKILESH', '6', 'B', '89'],
    ['SRI DHARSHINI ', '7', 'A', '99'],
    ['GOPIKA', '7', 'A', '90'],
    ['CHANDANA', '7', 'A', '90'],
    ['KARTHICK', '8', 'A', '95'],
    ['KALEESHWARAN', '8', 'B', '99'],
]
for row_idx in range(12):
    for col_idx in range(4):
        cell = table.cell(row_idx, col_idx)
        cell.text = text_content[row_idx][col_idx]
        if row_idx == 0:
            cell_xml = cell._element
            tc_pr = cell_xml.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), f'{background_rgb[0]:02X}{background_rgb[1]:02X}{background_rgb[2]:02X}')
            tc_pr.append(shading)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
for row in table.rows:
    row.cells[0].width = Pt(200)  # First column narrow
    row.cells[1].width = Pt(80)  # Middle column wide
    row.cells[2].width = Pt(80)
    row.cells[3].width = Pt(80)
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  
            border.set(qn('w:sz'), '4')  
            border.set(qn('w:space'), '0')  
            border.set(qn('w:color'), '000000')  
            borders.append(border)  
        tc_pr.append(borders)
doc.save("ppa_report.docx")
print("Table inserted")

#
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph('Linguistic Toppers:')
run = para.runs[0]
run.bold = True
run.font.size = Pt(12)
doc.save('ppa_report.docx')

# Table
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
doc = Document('ppa_report.docx')
table = doc.add_table(rows=10, cols=4)
background_rgb = RGBColor(255, 192, 0)
text_content = [
    ['Student Name ', 'Grade', 'Section', 'BSPI'],
    [Lingtoppers_name1, Lingtoppers_grade1, Lingtoppers_sec1, Lingtoppers_score1],
    [Lingtoppers_name2, Lingtoppers_grade2, Lingtoppers_sec2, Lingtoppers_score2],
    [Lingtoppers_name3, Lingtoppers_grade3, Lingtoppers_sec3, Lingtoppers_score3],
    [' ', ' ', ' ', ' '],
    ['RAHINI', '5', 'A', '67'],
    ['KARTHIK', '5', 'B', '60'],
    ['PRAGATHI', '6', 'A', '70'],
    ['VIDHARSHANA', '6', 'B', '72'],
    ['SAHAANA', '7', 'A', '98'],
]
for row_idx in range(10):
    for col_idx in range(4):
        cell = table.cell(row_idx, col_idx)
        cell.text = text_content[row_idx][col_idx]
        if row_idx == 0:
            cell_xml = cell._element
            tc_pr = cell_xml.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), f'{background_rgb[0]:02X}{background_rgb[1]:02X}{background_rgb[2]:02X}')
            tc_pr.append(shading)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
for row in table.rows:
    row.cells[0].width = Pt(200)  
    row.cells[1].width = Pt(80)  
    row.cells[2].width = Pt(80)
    row.cells[3].width = Pt(80)
for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._element
        tc_pr = cell_xml.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single') 
            border.set(qn('w:sz'), '4')  
            border.set(qn('w:space'), '0') 
            border.set(qn('w:color'), '000000')  
            borders.append(border)
        tc_pr.append(borders)
doc.save('ppa_report.docx')
print("Table created successfully")

# logo
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document('ppa_report.docx')  
para = doc.add_paragraph()
run = para.add_run()
logo_path = r'C:\Users\varsh\OneDrive\Sathyabama\Edsix\New folder\logo6.png'  
run.add_picture(logo_path, width=Inches(7)) 
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save('ppa_report.docx')  

# text center
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph()
para = doc.add_paragraph("Play is crucial to the healthy development of children's minds. Puzzles and strategy-based games help reinforce critical thinking skills.")
para.alignment = WD_ALIGN_PARAGRAPH.CENTER  
run = para.runs[0]
run.bold = True
run.font.size = Pt(12)  
run.font.color.rgb = RGBColor(164, 164, 164)
doc.save('ppa_report.docx')

#
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('8 Final Interface')
run = para.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(46, 116, 181)
doc.save('ppa_report.docx')

#
from docx import Document 
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph("Upon closely looking at the analysis and all the revealed data points, it seems reasonable to conclude that there is a visible development in the performance of each student who has been enrolled in the program. Nevertheless, it's vital to acknowledge the fact that certain skill areas come more naturally to students than others.")
para = doc.add_paragraph("Memory and Linguistics – seem to be the easier zones for students, while Problem Solving & Visual Processing are still a bit difficult to master.")
para = doc.add_paragraph("Specifically, the low scores in Visual processing stand out as a recurring challenge; this necessitates additional support for students to navigate and achieve better proficiency in this specific skill domain. The improvement seen in this skill domain is quite marginal and is a point that needs to be taken note of by the school.")
para = doc.add_paragraph("Overall, there is a lot of scope for improvement in each of the five skill domains. In the current scenario, where the focus is on cognitive upskilling, it is necessary for the school to recognize the connection between a program like SKILLANGELS and the academic performance of every student.")
run = para.runs[0]
run.font.size = Pt(12)
doc.save('ppa_report.docx')

#
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('9 Next Level Coonsiderations ')
run = para.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(46, 116, 181)
doc.save('ppa_report.docx')

#
from docx import Document
from docx.shared import Pt
doc = Document('ppa_report.docx')
para = doc.add_paragraph("Though we do agree that every bit of improvement in student performance calls for celebration, it should also help stakeholders be far from complacent about the results achieved. Students' performance at  ")
run = para.add_run( schoolname )
run.bold = True
para.add_run(' on the SKILLANGELS CAUP platform has moved from  ')
run = para.add_run( schoolavgbspi )
run.bold = True
run.font.color.rgb = RGBColor(0, 175, 80)
para.add_run(' While this is good it is certainly not enough!')
para = doc.add_paragraph("The opportunity for improvement in the domains of Problem Solving and Visual Processing highlights areas where teachers and schools can focus their efforts, offering more training to empower students in these skills. ")
para = doc.add_paragraph("Considering the government's prioritization of skill development, it is imperative for every student to actively participate in a platform that guides them in the right direction. The contemporary student needs to focus on practical skills crucial for real-life situations, and this program serves as a vital framework for building the necessary foundation.")
para = doc.add_paragraph("We express strong confidence that the ongoing implementation of SKILLANGELS - CAUP will positively elevate the performance of each student in the years ahead. ")
para = doc.add_paragraph("Thank you for joining us on this impactful journey to empower students with these valuable skills.")
run = para.runs[0]
run.font.size = Pt(12)
doc.save('ppa_report.docx')
print("Text inserted")

#
from docx import Document
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph('10 Other Links:')
run = para.runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(46, 116, 181) 
doc.save('ppa_report.docx')

# LAST TEXT  IN CENTER
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
doc = Document('ppa_report.docx')
para = doc.add_paragraph("Annexure 1 contains the individual scores of each student.")
para.alignment = WD_ALIGN_PARAGRAPH.CENTER  
run = para.runs[0]
run.bold = True
run.font.size = Pt(14)  
run.font.color.rgb = RGBColor(46, 116, 181) 
doc.save('ppa_report.docx')

#Adding footer text
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import RGBColor
doc = Document('ppa_report.docx')
section = doc.sections[0]
footer = section.footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT 
run = footer_paragraph.add_run( schoolname)
run.bold = True
run = footer_paragraph.add_run('\n')
run.font.color.rgb = RGBColor(0, 0, 0)
run.bold = True
run = footer_paragraph.add_run('Performance Report 2023-24')
run.font.color.rgb = RGBColor(0, 0, 0)
run.bold = True
footer_paragraph2 = footer.add_paragraph()
footer_paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT   
run = footer_paragraph2.add_run('Page | ')
run.font.color.rgb = RGBColor(0, 0, 0)
run.bold = True
field_code = 'PAGE'
run = footer_paragraph2.add_run()
field = OxmlElement('w:fldSimple')
field.set(qn('w:instr'), field_code)
run._r.append(field) 
doc.save('ppa_report.docx')
print("The pages are created sucessfully!!")

